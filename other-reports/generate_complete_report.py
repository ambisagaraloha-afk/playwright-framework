from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def add_heading_with_style(doc, text, level, color_rgb=(0, 51, 102)):
    """Add heading with custom styling"""
    heading = doc.add_heading(text, level=level)
    heading.runs[0].font.color.rgb = RGBColor(*color_rgb)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_colored_paragraph(doc, text, color_rgb=(0, 0, 0), bold=False, size=11):
    """Add paragraph with custom color"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(*color_rgb)
    run.font.bold = bold
    run.font.size = Pt(size)
    return p

def shade_cell(cell, fill_color):
    """Shade a table cell with color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_complete_report():
    """Create a complete, generic Playwright framework guide"""
    
    doc = Document()
    
    # Set up document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # ===== TITLE PAGE =====
    title = doc.add_heading('Complete Guide to E2E Test Automation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph('Playwright + TypeScript Framework Setup & Execution')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_paragraph()  # Spacing
    doc.add_paragraph()
    
    # Add subtitle info
    info_para = doc.add_paragraph('A Complete Methodology for Manual Testing to Automated Test Execution')
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.runs[0].font.size = Pt(12)
    info_para.runs[0].italic = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Project Details
    details_table = doc.add_table(rows=5, cols=2)
    details_table.style = 'Light Grid Accent 1'
    
    details = [
        ('Document Type', 'Complete Setup & Automation Guide'),
        ('Audience', 'QA Engineers, Testers, Developers (Beginners to Intermediate)'),
        ('Framework', 'Playwright + TypeScript'),
        ('Created Date', datetime.now().strftime("%B %d, %Y")),
        ('Status', 'Ready to Use'),
    ]
    
    for i, (key, value) in enumerate(details):
        key_cell = details_table.rows[i].cells[0]
        val_cell = details_table.rows[i].cells[1]
        key_cell.text = key
        val_cell.text = value
        shade_cell(key_cell, "4472C4")
        key_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    doc.add_page_break()
    
    # ===== TABLE OF CONTENTS =====
    add_heading_with_style(doc, 'Table of Contents', 1)
    toc_items = [
        '1. Introduction & Overview',
        '2. What You Will Learn',
        '3. Prerequisites & System Requirements',
        '4. Phase 1: Environment Setup',
        '   4.1 Installing Node.js',
        '   4.2 Installing Visual Studio Code',
        '   4.3 Installing VS Code Extensions',
        '   4.4 Creating Project Directory',
        '5. Phase 2: Framework Initialization',
        '   5.1 Initialize Node Project',
        '   5.2 Install Dependencies',
        '   5.3 Configure Playwright',
        '6. Phase 3: Project Structure Setup',
        '7. Phase 4: Manual Test Case Creation',
        '   7.1 Test Case Format & Standards',
        '   7.2 Test Case Template',
        '   7.3 Creating Excel Test Cases',
        '   7.4 Example Test Cases',
        '8. Phase 5: Automation Script Development',
        '   8.1 Page Object Model (POM) Pattern',
        '   8.2 Creating Page Classes',
        '   8.3 Creating Test Specifications',
        '   8.4 Test Data Management',
        '9. Phase 6: Test Execution',
        '   9.1 Running Tests',
        '   9.2 Viewing Test Reports',
        '10. Complete Workflow Example',
        '11. Best Practices & Troubleshooting',
        '12. Quick Reference Guide',
    ]
    
    for item in toc_items:
        indent_level = item.count('   ')
        clean_item = item.lstrip()
        if indent_level == 0:
            doc.add_paragraph(clean_item, style='List Bullet')
        else:
            doc.add_paragraph(clean_item, style='List Bullet 2')
    
    doc.add_page_break()
    
    # ===== 1. INTRODUCTION & OVERVIEW =====
    add_heading_with_style(doc, '1. Introduction & Overview', 1)
    doc.add_paragraph(
        'This guide provides a complete, step-by-step methodology for setting up and using the Playwright framework '
        'for end-to-end test automation. It covers everything from initial environment setup to creating and executing automated test cases.'
    )
    
    add_heading_with_style(doc, '1.1 What is Playwright?', 2)
    doc.add_paragraph(
        'Playwright is a modern, open-source testing framework that allows you to automate web applications across '
        'multiple browsers (Chromium, Firefox, WebKit). It supports multiple programming languages including TypeScript, JavaScript, Python, C#, and Java.'
    )
    
    features = [
        'Cross-browser testing (Chrome, Firefox, Safari)',
        'Multiple language support',
        'Powerful locator strategies',
        'Built-in test debugging tools',
        'Comprehensive reporting',
        'CI/CD integration ready',
        'Fast, reliable test execution',
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    add_heading_with_style(doc, '1.2 Why TypeScript?', 2)
    doc.add_paragraph(
        'TypeScript adds type safety to JavaScript, making code more robust and maintainable. '
        'It catches errors at compile-time rather than runtime, resulting in better code quality.'
    )
    
    ts_benefits = [
        'Type safety catches errors early',
        'Better IDE support and autocomplete',
        'Improved code documentation',
        'Easier maintenance and refactoring',
        'Better for large projects',
    ]
    for benefit in ts_benefits:
        doc.add_paragraph(benefit, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 2. WHAT YOU WILL LEARN =====
    add_heading_with_style(doc, '2. What You Will Learn', 1)
    doc.add_paragraph(
        'By following this guide, you will be able to:'
    )
    
    learning_outcomes = [
        'Set up a complete Playwright testing environment from scratch',
        'Understand and implement the Page Object Model (POM) design pattern',
        'Create well-structured manual test cases',
        'Convert manual test cases into automated scripts',
        'Execute tests across multiple browsers',
        'Generate and analyze test reports',
        'Debug and troubleshoot test failures',
        'Integrate tests with CI/CD pipelines',
        'Follow industry best practices for test automation',
    ]
    for outcome in learning_outcomes:
        doc.add_paragraph(outcome, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 3. PREREQUISITES & SYSTEM REQUIREMENTS =====
    add_heading_with_style(doc, '3. Prerequisites & System Requirements', 1)
    
    add_heading_with_style(doc, '3.1 System Requirements', 2)
    
    req_table = doc.add_table(rows=5, cols=2)
    req_table.style = 'Light Grid Accent 1'
    
    header_cells = req_table.rows[0].cells
    header_cells[0].text = 'Requirement'
    header_cells[1].text = 'Specification'
    shade_cell(header_cells[0], "4472C4")
    shade_cell(header_cells[1], "4472C4")
    header_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    header_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    reqs = [
        ('Operating System', 'Windows 10/11, macOS 10.15+, or Linux'),
        ('RAM', 'Minimum 4GB (8GB recommended)'),
        ('Disk Space', 'Minimum 2GB free space'),
        ('Internet Connection', 'Required for downloading packages and running tests'),
    ]
    
    for i, (req, spec) in enumerate(reqs):
        row = req_table.rows[i + 1]
        row.cells[0].text = req
        row.cells[1].text = spec
    
    add_heading_with_style(doc, '3.2 Software to Install', 2)
    
    software_table = doc.add_table(rows=6, cols=3)
    software_table.style = 'Light Grid Accent 1'
    
    header_cells = software_table.rows[0].cells
    headers = ['Software', 'Version', 'Download Link']
    for i, header_text in enumerate(headers):
        header_cells[i].text = header_text
        shade_cell(header_cells[i], "4472C4")
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    software = [
        ('Node.js', '18.x LTS or higher', 'https://nodejs.org/'),
        ('Visual Studio Code', 'Latest version', 'https://code.visualstudio.com/'),
        ('Git (Optional)', 'Latest version', 'https://git-scm.com/'),
        ('Playwright', 'Installed via npm', 'npm install @playwright/test'),
        ('TypeScript', 'Installed via npm', 'npm install typescript'),
    ]
    
    for i, (soft, version, link) in enumerate(software):
        row = software_table.rows[i + 1]
        row.cells[0].text = soft
        row.cells[1].text = version
        row.cells[2].text = link
    
    doc.add_page_break()
    
    # ===== 4. PHASE 1: ENVIRONMENT SETUP =====
    add_heading_with_style(doc, '4. Phase 1: Environment Setup', 1)
    doc.add_paragraph(
        'In this phase, we will install all necessary software and tools to prepare your system for test automation development.'
    )
    
    # 4.1 Installing Node.js
    add_heading_with_style(doc, '4.1 Installing Node.js', 2)
    doc.add_paragraph(
        'Node.js is a JavaScript runtime that allows us to run JavaScript code outside the browser. '
        'It\'s the foundation for running Playwright and managing project dependencies.'
    )
    
    add_colored_paragraph(doc, 'Step 1: Download Node.js', bold=True, color_rgb=(0, 102, 204))
    steps_1_1 = [
        'Go to https://nodejs.org/',
        'Click on "LTS" (Long Term Support) version - this is the recommended stable version',
        'The download will start automatically for your operating system',
    ]
    for step in steps_1_1:
        doc.add_paragraph(step, style='List Number')
    
    add_colored_paragraph(doc, 'Step 2: Install Node.js', bold=True, color_rgb=(0, 102, 204))
    steps_1_2 = [
        'Open the downloaded installer file',
        'Click "Next" and accept the license agreement',
        'Choose installation path (default location is fine)',
        'In "Tools for Native Modules", check the checkbox for optional build tools',
        'Click "Install" and wait for completion',
        'Click "Finish" when installation is complete',
    ]
    for step in steps_1_2:
        doc.add_paragraph(step, style='List Number')
    
    add_colored_paragraph(doc, 'Step 3: Verify Installation', bold=True, color_rgb=(0, 102, 204))
    doc.add_paragraph('Open Command Prompt or PowerShell and run:', style='List Number')
    verify_cmds = [
        'node --version',
        'npm --version',
    ]
    for cmd in verify_cmds:
        cmd_para = doc.add_paragraph(cmd)
        for run in cmd_para.runs:
            run.font.name = 'Courier New'
            run.font.bold = True
    
    doc.add_paragraph('You should see version numbers displayed (e.g., v18.17.0)')
    
    # 4.2 Installing Visual Studio Code
    add_heading_with_style(doc, '4.2 Installing Visual Studio Code', 2)
    doc.add_paragraph(
        'Visual Studio Code is a lightweight but powerful source code editor with excellent support '
        'for JavaScript, TypeScript, and extensions for test automation.'
    )
    
    add_colored_paragraph(doc, 'Step 1: Download VS Code', bold=True, color_rgb=(0, 102, 204))
    steps_2_1 = [
        'Go to https://code.visualstudio.com/',
        'Click the download button for your operating system',
        'The appropriate installer will download',
    ]
    for step in steps_2_1:
        doc.add_paragraph(step, style='List Number')
    
    add_colored_paragraph(doc, 'Step 2: Install VS Code', bold=True, color_rgb=(0, 102, 204))
    steps_2_2 = [
        'Run the downloaded installer',
        'Accept the license agreement',
        'Choose installation path',
        'Select additional tasks (recommended: Add to PATH, Register as default editor)',
        'Click "Install" and wait for completion',
    ]
    for step in steps_2_2:
        doc.add_paragraph(step, style='List Number')
    
    add_colored_paragraph(doc, 'Step 3: Launch VS Code', bold=True, color_rgb=(0, 102, 204))
    doc.add_paragraph('Double-click the VS Code icon on your desktop to launch the editor.')
    
    # 4.3 Installing VS Code Extensions
    add_heading_with_style(doc, '4.3 Installing VS Code Extensions', 2)
    doc.add_paragraph(
        'Extensions enhance VS Code functionality. Install these essential extensions for test automation development:'
    )
    
    add_colored_paragraph(doc, 'Required Extensions:', bold=True, color_rgb=(0, 102, 204))
    
    extensions = {
        'Playwright Test for VS Code': 'Run and debug Playwright tests directly from the editor',
        'GitHub Copilot': 'AI-powered code assistance (optional but very helpful)',
        'ESLint': 'JavaScript linting and code quality checks',
        'Prettier - Code formatter': 'Automatic code formatting',
        'Thunder Client': 'API testing tool (optional)',
        'REST Client': 'Make HTTP requests from VS Code',
    }
    
    for ext_name, ext_purpose in extensions.items():
        doc.add_paragraph(f'{ext_name}: {ext_purpose}', style='List Bullet')
    
    add_colored_paragraph(doc, 'How to Install Extensions:', bold=True, color_rgb=(0, 102, 204))
    ext_steps = [
        'Click the Extensions icon in the left sidebar (looks like four squares)',
        'Search for the extension name in the search box',
        'Click "Install" button',
        'Wait for the installation to complete',
        'Repeat for each extension',
    ]
    for step in ext_steps:
        doc.add_paragraph(step, style='List Number')
    
    # 4.4 Creating Project Directory
    add_heading_with_style(doc, '4.4 Creating Project Directory', 2)
    
    add_colored_paragraph(doc, 'Step 1: Create a New Folder', bold=True, color_rgb=(0, 102, 204))
    doc.add_paragraph('Create a new folder on your system (e.g., C:\\Users\\YourName\\Documents\\Playwright_Framework)')
    
    add_colored_paragraph(doc, 'Step 2: Open in VS Code', bold=True, color_rgb=(0, 102, 204))
    steps_open = [
        'Open VS Code',
        'Click "File" > "Open Folder"',
        'Navigate to and select your Playwright_Framework folder',
        'Click "Select Folder"',
    ]
    for step in steps_open:
        doc.add_paragraph(step, style='List Number')
    
    add_colored_paragraph(doc, 'Step 3: Open Terminal in VS Code', bold=True, color_rgb=(0, 102, 204))
    doc.add_paragraph('Press Ctrl + ` (backtick) or go to Terminal > New Terminal')
    
    doc.add_page_break()
    
    # ===== 5. PHASE 2: FRAMEWORK INITIALIZATION =====
    add_heading_with_style(doc, '5. Phase 2: Framework Initialization', 1)
    doc.add_paragraph(
        'In this phase, we initialize the Node.js project and install all required dependencies.'
    )
    
    # 5.1 Initialize Node Project
    add_heading_with_style(doc, '5.1 Initialize Node Project', 2)
    
    add_colored_paragraph(doc, 'Command:', bold=True, color_rgb=(0, 102, 204))
    cmd_para = doc.add_paragraph('npm init -y')
    for run in cmd_para.runs:
        run.font.name = 'Courier New'
        run.font.bold = True
        run.font.size = Pt(11)
    
    doc.add_paragraph('This command creates a package.json file with default settings. The -y flag skips interactive prompts.')
    
    add_colored_paragraph(doc, 'What is package.json?', bold=True, color_rgb=(0, 102, 204))
    doc.add_paragraph(
        'package.json is a configuration file that stores information about your project and its dependencies. '
        'It\'s used by npm to manage packages and project metadata.'
    )
    
    # 5.2 Install Dependencies
    add_heading_with_style(doc, '5.2 Install Dependencies', 2)
    
    add_colored_paragraph(doc, 'Install Playwright:', bold=True, color_rgb=(0, 102, 204))
    cmd_para = doc.add_paragraph('npm install --save-dev @playwright/test')
    for run in cmd_para.runs:
        run.font.name = 'Courier New'
        run.font.bold = True
    
    add_colored_paragraph(doc, 'Install TypeScript Types:', bold=True, color_rgb=(0, 102, 204))
    cmd_para = doc.add_paragraph('npm install --save-dev @types/node')
    for run in cmd_para.runs:
        run.font.name = 'Courier New'
        run.font.bold = True
    
    add_colored_paragraph(doc, 'Install Excel Parser (for reading test cases from Excel):', bold=True, color_rgb=(0, 102, 204))
    cmd_para = doc.add_paragraph('npm install xlsx')
    for run in cmd_para.runs:
        run.font.name = 'Courier New'
        run.font.bold = True
    
    # 5.3 Configure Playwright
    add_heading_with_style(doc, '5.3 Configure Playwright', 2)
    
    add_colored_paragraph(doc, 'Initialize Playwright Config:', bold=True, color_rgb=(0, 102, 204))
    cmd_para = doc.add_paragraph('npx playwright install')
    for run in cmd_para.runs:
        run.font.name = 'Courier New'
        run.font.bold = True
    
    doc.add_paragraph('This downloads browser binaries for Chromium, Firefox, and WebKit.')
    
    add_colored_paragraph(doc, 'Create playwright.config.ts:', bold=True, color_rgb=(0, 102, 204))
    doc.add_paragraph('Create a new file named "playwright.config.ts" in your project root with the following content:')
    
    config_code = '''import { defineConfig, devices } from '@playwright/test';

export default defineConfig({
  testDir: './tests',
  fullyParallel: true,
  forbidOnly: !!process.env.CI,
  retries: process.env.CI ? 2 : 0,
  workers: process.env.CI ? 1 : undefined,
  reporter: 'html',
  timeout: 60000,
  
  use: {
    trace: 'on-first-retry',
    navigationTimeout: 30000,
    screenshot: 'only-on-failure',
    video: 'retain-on-failure',
  },

  projects: [
    {
      name: 'chromium',
      use: { ...devices['Desktop Chrome'] },
    },
    {
      name: 'firefox',
      use: { ...devices['Desktop Firefox'] },
    },
    {
      name: 'webkit',
      use: { ...devices['Desktop Safari'] },
    },
  ],
});'''
    
    code_para = doc.add_paragraph(config_code, style='No Spacing')
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # ===== 6. PHASE 3: PROJECT STRUCTURE SETUP =====
    add_heading_with_style(doc, '6. Phase 3: Project Structure Setup', 1)
    doc.add_paragraph(
        'Organize your project with a clear folder structure. This makes your project maintainable and scalable.'
    )
    
    add_heading_with_style(doc, '6.1 Creating Folder Structure', 2)
    
    add_colored_paragraph(doc, 'Create the following folders in your project root:', bold=True, color_rgb=(0, 102, 204))
    
    folders_to_create = [
        'Pages - For Page Object Model classes',
        'tests - For test specification files',
        'utilities - For helper functions and test data',
        'test-cases_manual - For storing manual test cases',
        'test-results - For test execution logs',
    ]
    
    for folder in folders_to_create:
        doc.add_paragraph(folder, style='List Bullet')
    
    add_heading_with_style(doc, '6.2 Final Project Structure', 2)
    
    structure_text = '''Playwright_Framework/
├── Pages/
│   ├── basePage.ts              # Base class for all pages
│   ├── homePage.ts              # Home page objects and methods
│   ├── loginPage.ts             # Login page objects and methods
│   └── cartPage.ts              # Shopping cart objects and methods
│
├── tests/
│   ├── auth.spec.ts             # Authentication tests
│   ├── shopping.spec.ts         # Shopping cart tests
│   ├── checkout.spec.ts         # Checkout process tests
│   └── regression.spec.ts       # Regression test suite
│
├── utilities/
│   ├── testData.ts              # Test data constants
│   ├── helpers.ts               # Helper functions
│   └── logger.ts                # Logging utility
│
├── test-cases_manual/
│   ├── TestCases.xlsx           # Manual test cases
│   └── test_documentation.md    # Test documentation
│
├── playwright.config.ts         # Playwright configuration
├── package.json                 # Project dependencies
├── tsconfig.json                # TypeScript configuration
└── README.md                    # Project documentation'''
    
    struct_para = doc.add_paragraph(structure_text, style='No Spacing')
    for run in struct_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # ===== 7. PHASE 4: MANUAL TEST CASE CREATION =====
    add_heading_with_style(doc, '7. Phase 4: Manual Test Case Creation', 1)
    doc.add_paragraph(
        'Before automating tests, we create comprehensive manual test cases. These serve as blueprints for automation.'
    )
    
    # 7.1 Test Case Format & Standards
    add_heading_with_style(doc, '7.1 Test Case Format & Standards', 2)
    
    add_colored_paragraph(doc, 'Why Manual Test Cases First?', bold=True, color_rgb=(0, 102, 204))
    doc.add_paragraph(
        'Manual test cases serve multiple purposes: they document expected application behavior, '
        'serve as requirements for automation, and help identify gaps in testing strategy.'
    )
    
    add_heading_with_style(doc, '7.2 Test Case Template', 2)
    
    template_table = doc.add_table(rows=12, cols=2)
    template_table.style = 'Light Grid Accent 1'
    
    template_data = [
        ('Test Case ID', 'Unique identifier (e.g., TC-001)'),
        ('Test Case Name', 'Brief descriptive name'),
        ('Module/Feature', 'Which feature is being tested'),
        ('Description', 'Detailed description of test objective'),
        ('Priority', 'High / Medium / Low'),
        ('Severity', 'Critical / Major / Minor'),
        ('Pre-conditions', 'Setup requirements before test'),
        ('Test Data', 'Input values and test data'),
        ('Steps', 'Numbered steps to execute'),
        ('Expected Result', 'Expected outcome or assertion'),
        ('Post-conditions', 'Cleanup after test execution'),
        ('Status', 'Pass / Fail / Not Executed'),
    ]
    
    for i, (field, description) in enumerate(template_data):
        row = template_table.rows[i]
        row.cells[0].text = field
        row.cells[1].text = description
        shade_cell(row.cells[0], "E7E6E6")
    
    # 7.3 Creating Excel Test Cases
    add_heading_with_style(doc, '7.3 Creating Excel Test Cases', 2)
    
    add_colored_paragraph(doc, 'Step 1: Create Excel File', bold=True, color_rgb=(0, 102, 204))
    excel_steps_1 = [
        'Open Microsoft Excel or Google Sheets',
        'Create a new spreadsheet',
        'Save as "TestCases.xlsx" in the test-cases_manual folder',
    ]
    for step in excel_steps_1:
        doc.add_paragraph(step, style='List Number')
    
    add_colored_paragraph(doc, 'Step 2: Create Column Headers', bold=True, color_rgb=(0, 102, 204))
    doc.add_paragraph('In the first row, add these column headers:')
    
    headers_list = [
        'A: TestCaseID',
        'B: TestCaseName',
        'C: Module',
        'D: Description',
        'E: Priority',
        'F: Severity',
        'G: Preconditions',
        'H: TestData',
        'I: Steps',
        'J: ExpectedResult',
        'K: PostConditions',
        'L: Status',
    ]
    
    for header in headers_list:
        doc.add_paragraph(header, style='List Bullet')
    
    add_colored_paragraph(doc, 'Step 3: Add Test Cases', bold=True, color_rgb=(0, 102, 204))
    doc.add_paragraph('Fill in the rows with your test case data. Save the file regularly.')
    
    # 7.4 Example Test Cases
    add_heading_with_style(doc, '7.4 Example Test Cases', 2)
    
    add_colored_paragraph(doc, 'Example 1: User Login Test Case', bold=True, color_rgb=(0, 102, 204))
    
    example1_table = doc.add_table(rows=11, cols=2)
    example1_table.style = 'Light Grid'
    
    example1_data = [
        ('TestCaseID', 'TC-001'),
        ('TestCaseName', 'Verify user can login with valid credentials'),
        ('Module', 'Authentication'),
        ('Priority', 'High'),
        ('Preconditions', 'User account exists, User not logged in'),
        ('TestData', 'Username: testuser@example.com, Password: Test@123'),
        ('Steps', '1. Navigate to login page\n2. Enter username\n3. Enter password\n4. Click Login button'),
        ('ExpectedResult', 'User is logged in successfully, Dashboard is displayed'),
        ('Status', 'Not Executed'),
    ]
    
    for i, (key, value) in enumerate(example1_data):
        row = example1_table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
        shade_cell(row.cells[0], "D9E1F2")
    
    doc.add_page_break()
    
    add_colored_paragraph(doc, 'Example 2: Add to Cart Test Case', bold=True, color_rgb=(0, 102, 204))
    
    example2_table = doc.add_table(rows=10, cols=2)
    example2_table.style = 'Light Grid'
    
    example2_data = [
        ('TestCaseID', 'TC-002'),
        ('TestCaseName', 'Verify user can add product to cart'),
        ('Module', 'Shopping Cart'),
        ('Priority', 'High'),
        ('Preconditions', 'User is on product listing page'),
        ('TestData', 'Product: Samsung Galaxy S10'),
        ('Steps', '1. Locate product\n2. Click Add to Cart button\n3. Verify product in cart'),
        ('ExpectedResult', 'Product is added to cart, Cart count updated'),
        ('Status', 'Not Executed'),
    ]
    
    for i, (key, value) in enumerate(example2_data):
        row = example2_table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
        shade_cell(row.cells[0], "D9E1F2")
    
    doc.add_page_break()
    
    # ===== 8. PHASE 5: AUTOMATION SCRIPT DEVELOPMENT =====
    add_heading_with_style(doc, '8. Phase 5: Automation Script Development', 1)
    doc.add_paragraph(
        'In this phase, we convert manual test cases into automated scripts using the Page Object Model pattern.'
    )
    
    # 8.1 Page Object Model Pattern
    add_heading_with_style(doc, '8.1 Page Object Model (POM) Pattern', 2)
    
    doc.add_paragraph(
        'Page Object Model is a design pattern that enhances test maintenance and reduces code duplication. '
        'Each page is represented as a class with:',
        style='List Bullet'
    )
    
    pom_components = [
        'Page elements (locators) as properties',
        'Actions on those elements as methods',
        'Separation of test logic from UI locators',
    ]
    
    for component in pom_components:
        doc.add_paragraph(component, style='List Bullet')
    
    add_colored_paragraph(doc, 'Benefits of POM:', bold=True, color_rgb=(0, 102, 204))
    pom_benefits = [
        'Easy maintenance - Update locators in one place',
        'Reduced code duplication',
        'Better test readability',
        'Scalable test framework',
        'Reusable page components',
    ]
    for benefit in pom_benefits:
        doc.add_paragraph(benefit, style='List Bullet')
    
    # 8.2 Creating Page Classes
    add_heading_with_style(doc, '8.2 Creating Page Classes', 2)
    
    add_colored_paragraph(doc, 'Example: Login Page Class', bold=True, color_rgb=(0, 102, 204))
    
    login_page_code = '''import { Page, Locator } from '@playwright/test';

export class LoginPage {
    readonly page: Page;
    readonly emailInput: Locator;
    readonly passwordInput: Locator;
    readonly loginButton: Locator;
    readonly errorMessage: Locator;

    constructor(page: Page) {
        this.page = page;
        // Define locators using XPath or CSS selectors
        this.emailInput = page.locator('input[id="loginusername"]');
        this.passwordInput = page.locator('input[id="loginpassword"]');
        this.loginButton = page.locator('button:has-text("Log in")');
        this.errorMessage = page.locator('.alert-danger');
    }

    // Method to navigate to login page
    async navigateToLoginPage() {
        await this.page.goto('https://demoblaze.com/index.html');
        await this.page.locator('#login2').click();
    }

    // Method to enter email
    async enterEmail(email: string) {
        await this.emailInput.fill(email);
    }

    // Method to enter password
    async enterPassword(password: string) {
        await this.passwordInput.fill(password);
    }

    // Method to click login button
    async clickLoginButton() {
        await this.loginButton.click();
    }

    // Method to perform complete login
    async login(email: string, password: string) {
        await this.navigateToLoginPage();
        await this.enterEmail(email);
        await this.enterPassword(password);
        await this.clickLoginButton();
    }

    // Method to verify login error
    async getErrorMessage() {
        return await this.errorMessage.textContent();
    }
}'''
    
    code_para = doc.add_paragraph(login_page_code, style='No Spacing')
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
    
    doc.add_page_break()
    
    # 8.3 Creating Test Specifications
    add_heading_with_style(doc, '8.3 Creating Test Specifications', 2)
    
    add_colored_paragraph(doc, 'Test Specification File: tests/auth.spec.ts', bold=True, color_rgb=(0, 102, 204))
    
    test_spec_code = '''import { test, expect } from '@playwright/test';
import { LoginPage } from '../Pages/loginPage';
import { credentials } from '../utilities/testData';

test.describe('Authentication Tests', () => {
    let loginPage: LoginPage;

    test.beforeEach(async ({ page }) => {
        loginPage = new LoginPage(page);
    });

    test('TC-001: User should login with valid credentials', 
        async ({ page }) => {
            // Navigate to login
            await loginPage.login(credentials.username, credentials.password);
            
            // Verify user is logged in
            const welcomeMessage = page.locator('.navbar-text');
            await expect(welcomeMessage).toContainText(credentials.username);
        }
    );

    test('TC-003: User should see error with invalid credentials', 
        async ({ page }) => {
            // Try to login with wrong password
            await loginPage.login('testuser@example.com', 'wrongpassword');
            
            // Verify error message appears
            const errorMsg = await loginPage.getErrorMessage();
            expect(errorMsg).toContain('User does not exist');
        }
    );

    test('TC-004: User should logout successfully', 
        async ({ page }) => {
            // Login first
            await loginPage.login(credentials.username, credentials.password);
            
            // Click logout
            await page.locator('#logout2').click();
            
            // Verify login button is visible again
            const loginButton = page.locator('#login2');
            await expect(loginButton).toBeVisible();
        }
    );
});'''
    
    code_para = doc.add_paragraph(test_spec_code, style='No Spacing')
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
    
    # 8.4 Test Data Management
    add_heading_with_style(doc, '8.4 Test Data Management', 2)
    
    doc.add_paragraph('Create a utilities/testData.ts file to store test data:')
    
    test_data_code = '''export const credentials = {
    username: 'testuser@example.com',
    password: 'Test@123',
    invalidPassword: 'wrongpassword',
};

export const testProducts = {
    phone: 'Samsung Galaxy S10',
    laptop: 'Dell Inspiron 15',
    monitor: 'LG 24 inch',
};

export const testUsers = {
    validEmail: 'valid@example.com',
    invalidEmail: 'invalid@example.com',
    existingUser: 'existinguser@example.com',
};

export const applicationUrls = {
    baseUrl: 'https://demoblaze.com',
    loginUrl: 'https://demoblaze.com/index.html',
};'''
    
    code_para = doc.add_paragraph(test_data_code, style='No Spacing')
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # ===== 9. PHASE 6: TEST EXECUTION =====
    add_heading_with_style(doc, '9. Phase 6: Test Execution', 1)
    
    # 9.1 Running Tests
    add_heading_with_style(doc, '9.1 Running Tests', 2)
    
    add_heading_with_style(doc, '9.1.1 Run All Tests', 3)
    cmd_para = doc.add_paragraph('npx playwright test')
    for run in cmd_para.runs:
        run.font.name = 'Courier New'
        run.font.bold = True
    
    add_heading_with_style(doc, '9.1.2 Run Specific Test File', 3)
    cmd_para = doc.add_paragraph('npx playwright test tests/auth.spec.ts')
    for run in cmd_para.runs:
        run.font.name = 'Courier New'
        run.font.bold = True
    
    add_heading_with_style(doc, '9.1.3 Run Specific Test Case', 3)
    cmd_para = doc.add_paragraph('npx playwright test -g "TC-001"')
    for run in cmd_para.runs:
        run.font.name = 'Courier New'
        run.font.bold = True
    
    add_heading_with_style(doc, '9.1.4 Run in Headed Mode (see browser)', 3)
    cmd_para = doc.add_paragraph('npx playwright test --headed')
    for run in cmd_para.runs:
        run.font.name = 'Courier New'
        run.font.bold = True
    
    add_heading_with_style(doc, '9.1.5 Run in Debug Mode', 3)
    cmd_para = doc.add_paragraph('npx playwright test --debug')
    for run in cmd_para.runs:
        run.font.name = 'Courier New'
        run.font.bold = True
    
    add_heading_with_style(doc, '9.1.6 Run Specific Browser', 3)
    cmd_para = doc.add_paragraph('npx playwright test --project=chromium')
    for run in cmd_para.runs:
        run.font.name = 'Courier New'
        run.font.bold = True
    doc.add_paragraph('Options: chromium, firefox, webkit')
    
    add_heading_with_style(doc, '9.1.7 Run with Verbose Output', 3)
    cmd_para = doc.add_paragraph('npx playwright test --verbose')
    for run in cmd_para.runs:
        run.font.name = 'Courier New'
        run.font.bold = True
    
    # 9.2 Viewing Test Reports
    add_heading_with_style(doc, '9.2 Viewing Test Reports', 2)
    
    doc.add_paragraph(
        'After test execution, Playwright generates an HTML report automatically. '
        'To view the report:'
    )
    
    report_steps = [
        'Run tests with: npx playwright test',
        'View report with: npx playwright show-report',
        'This opens the report in your default browser',
        'The report shows: Test results, screenshots, videos, traces, timing details',
    ]
    
    for i, step in enumerate(report_steps):
        doc.add_paragraph(step, style='List Number')
    
    add_colored_paragraph(doc, 'Report Location:', bold=True, color_rgb=(0, 102, 204))
    doc.add_paragraph('The HTML report is stored in the playwright-report/ folder')
    
    doc.add_page_break()
    
    # ===== 10. COMPLETE WORKFLOW EXAMPLE =====
    add_heading_with_style(doc, '10. Complete Workflow Example', 1)
    doc.add_paragraph(
        'This section shows a complete example workflow from manual test case to automated execution.'
    )
    
    add_heading_with_style(doc, '10.1 Scenario: E-commerce Product Search', 2)
    
    add_colored_paragraph(doc, 'Step 1: Create Manual Test Case', bold=True, color_rgb=(0, 102, 204))
    
    workflow_example = '''
Manual Test Case: TC-005 - Search Products by Category

TestCaseID: TC-005
TestCaseName: User should filter products by category
Module: Product Browsing
Priority: High
Preconditions:
  - User is on the home page
  - Product categories are available (Phones, Laptops, Monitors)
TestData:
  - Category: Laptops
Steps:
  1. Navigate to https://demoblaze.com
  2. Wait for page to load
  3. Click on "Laptops" category
  4. Verify laptop products are displayed
  5. Verify product count matches category
ExpectedResult:
  - Laptops category is selected (highlighted)
  - Only laptop products are displayed
  - Product count shows multiple laptops
Status: Not Executed
    '''
    
    workflow_para = doc.add_paragraph(workflow_example, style='No Spacing')
    for run in workflow_para.runs:
        run.font.size = Pt(10)
    
    add_colored_paragraph(doc, 'Step 2: Create Page Class', bold=True, color_rgb=(0, 102, 204))
    
    page_class_example = '''// Pages/productPage.ts
export class ProductPage {
    readonly page: Page;
    readonly laptopsCategoryLink: Locator;
    readonly productItems: Locator;

    constructor(page: Page) {
        this.page = page;
        this.laptopsCategoryLink = page.locator('a:has-text("Laptops")');
        this.productItems = page.locator('div.col-lg-4');
    }

    async clickLaptopsCategory() {
        await this.laptopsCategoryLink.click();
        await this.page.waitForTimeout(1000);
    }

    async getProductCount() {
        return await this.productItems.count();
    }
}'''
    
    page_para = doc.add_paragraph(page_class_example, style='No Spacing')
    for run in page_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
    
    add_colored_paragraph(doc, 'Step 3: Create Test Specification', bold=True, color_rgb=(0, 102, 204))
    
    test_spec_example = '''// tests/productBrowsing.spec.ts
test('TC-005: Filter products by category', async ({ page }) => {
    const productPage = new ProductPage(page);
    
    // Navigate to home page
    await page.goto('https://demoblaze.com');
    
    // Click on Laptops category
    await productPage.clickLaptopsCategory();
    
    // Get product count
    const count = await productPage.getProductCount();
    
    // Verify at least one product is displayed
    expect(count).toBeGreaterThan(0);
});'''
    
    test_para = doc.add_paragraph(test_spec_example, style='No Spacing')
    for run in test_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
    
    add_colored_paragraph(doc, 'Step 4: Execute Test', bold=True, color_rgb=(0, 102, 204))
    
    execute_para = doc.add_paragraph('npx playwright test tests/productBrowsing.spec.ts --headed')
    for run in execute_para.runs:
        run.font.name = 'Courier New'
        run.font.bold = True
    
    add_colored_paragraph(doc, 'Step 5: View Results', bold=True, color_rgb=(0, 102, 204))
    
    results_para = doc.add_paragraph('npx playwright show-report')
    for run in results_para.runs:
        run.font.name = 'Courier New'
        run.font.bold = True
    
    doc.add_page_break()
    
    # ===== 11. BEST PRACTICES & TROUBLESHOOTING =====
    add_heading_with_style(doc, '11. Best Practices & Troubleshooting', 1)
    
    add_heading_with_style(doc, '11.1 Best Practices', 2)
    
    best_practices = {
        'Use Page Object Model': [
            'Separate UI locators from test logic',
            'Create reusable page methods',
            'Maintain a base page class for common functionality',
        ],
        'Explicit Waits': [
            'Use waitFor() for element visibility instead of Thread.sleep()',
            'Set appropriate timeouts based on application behavior',
            'Wait for specific element conditions',
        ],
        'Test Data': [
            'Keep test data separate from test code',
            'Use environment variables for sensitive data',
            'Create data factories for complex test scenarios',
        ],
        'Error Handling': [
            'Use try-catch blocks for critical operations',
            'Provide meaningful error messages',
            'Log important test steps for debugging',
        ],
        'Test Maintenance': [
            'Write descriptive test names',
            'Add comments for complex logic',
            'Review and update locators regularly',
        ],
    }
    
    for practice, tips in best_practices.items():
        add_heading_with_style(doc, practice, 3, (0, 102, 204))
        for tip in tips:
            doc.add_paragraph(tip, style='List Bullet')
    
    add_heading_with_style(doc, '11.2 Common Issues & Solutions', 2)
    
    issues_table = doc.add_table(rows=7, cols=2)
    issues_table.style = 'Light Grid Accent 1'
    
    header_cells = issues_table.rows[0].cells
    header_cells[0].text = 'Issue'
    header_cells[1].text = 'Solution'
    shade_cell(header_cells[0], "4472C4")
    shade_cell(header_cells[1], "4472C4")
    header_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    header_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    issues = [
        ('Element not found', 'Verify locator is correct, check page load timing, use waitFor()'),
        ('Test timeout', 'Increase timeout in config, add proper waits, check network'),
        ('Flaky tests', 'Use explicit waits, add retry logic, improve locators'),
        ('Browser crash', 'Update Playwright, check system resources, verify test cleanup'),
        ('Report not generating', 'Ensure HTML reporter is configured, check output path permissions'),
        ('Tests run too slow', 'Use parallelization, optimize waits, check test logic'),
    ]
    
    for i, (issue, solution) in enumerate(issues):
        row = issues_table.rows[i + 1]
        row.cells[0].text = issue
        row.cells[1].text = solution
    
    doc.add_page_break()
    
    # ===== 12. QUICK REFERENCE GUIDE =====
    add_heading_with_style(doc, '12. Quick Reference Guide', 1)
    
    add_heading_with_style(doc, '12.1 Essential Commands', 2)
    
    commands_table = doc.add_table(rows=12, cols=2)
    commands_table.style = 'Light Grid Accent 1'
    
    header_cells = commands_table.rows[0].cells
    header_cells[0].text = 'Command'
    header_cells[1].text = 'Description'
    shade_cell(header_cells[0], "4472C4")
    shade_cell(header_cells[1], "4472C4")
    header_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    header_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    commands = [
        ('npm init -y', 'Initialize Node.js project'),
        ('npm install --save-dev @playwright/test', 'Install Playwright'),
        ('npx playwright install', 'Download browser binaries'),
        ('npx playwright test', 'Run all tests'),
        ('npx playwright test tests/file.spec.ts', 'Run specific test file'),
        ('npx playwright test -g "test name"', 'Run tests matching pattern'),
        ('npx playwright test --headed', 'Run tests with visible browser'),
        ('npx playwright test --debug', 'Run with Playwright Inspector'),
        ('npx playwright show-report', 'View test report'),
        ('npx playwright codegen', 'Record user actions to generate code'),
        ('npm install xlsx', 'Install Excel parser'),
    ]
    
    for i, (cmd, desc) in enumerate(commands):
        row = commands_table.rows[i + 1]
        cmd_cell = row.cells[0]
        desc_cell = row.cells[1]
        cmd_cell.text = cmd
        desc_cell.text = desc
        
        for run in cmd_cell.paragraphs[0].runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
    
    add_heading_with_style(doc, '12.2 Common Locator Strategies', 2)
    
    locator_table = doc.add_table(rows=7, cols=2)
    locator_table.style = 'Light Grid Accent 1'
    
    header_cells = locator_table.rows[0].cells
    header_cells[0].text = 'Strategy'
    header_cells[1].text = 'Example'
    shade_cell(header_cells[0], "4472C4")
    shade_cell(header_cells[1], "4472C4")
    header_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    header_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    locators = [
        ('ID', 'page.locator("#elementId")'),
        ('Class', 'page.locator(".className")'),
        ('XPath', 'page.locator("//button[text()=\'Click\']")'),
        ('CSS Selector', 'page.locator("input[type=\'email\']")'),
        ('Text Content', 'page.locator("button:has-text(\'Login\')")'),
        ('Attribute', 'page.locator("[aria-label=\'Search\']")'),
    ]
    
    for i, (strategy, example) in enumerate(locators):
        row = locator_table.rows[i + 1]
        row.cells[0].text = strategy
        example_cell = row.cells[1]
        example_cell.text = example
        
        for run in example_cell.paragraphs[0].runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
    
    add_heading_with_style(doc, '12.3 Assertion Examples', 2)
    
    assertion_code = '''// Visibility assertions
await expect(element).toBeVisible();
await expect(element).toBeHidden();

// Content assertions
await expect(element).toContainText('Expected text');
await expect(element).toHaveText('Exact text');

// Value assertions
await expect(input).toHaveValue('expected value');

// Count assertions
await expect(elements).toHaveCount(5);

// Attribute assertions
await expect(element).toHaveAttribute('class', 'active');

// State assertions
await expect(element).toBeEnabled();
await expect(element).toBeDisabled();
await expect(element).toBeChecked();

// URL assertions
await expect(page).toHaveURL('https://example.com');
'''
    
    code_para = doc.add_paragraph(assertion_code, style='No Spacing')
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # ===== CONCLUSION =====
    add_heading_with_style(doc, 'Conclusion', 1)
    
    doc.add_paragraph(
        'You now have a complete understanding of how to set up and use a Playwright testing framework. '
        'The methodology presented in this guide ensures that your tests are maintainable, scalable, and efficient.'
    )
    
    doc.add_paragraph()
    
    key_takeaways = [
        'Follow a structured approach: Manual tests → Page Objects → Automation',
        'Use Page Object Model for maintainability',
        'Separate test logic from UI locators',
        'Implement explicit waits instead of hard-coded delays',
        'Maintain test data separately from test code',
        'Review and update locators regularly',
        'Use meaningful test names and documentation',
        'Implement proper error handling and logging',
        'Use version control (Git) for your code',
        'Continuously improve your test suite',
    ]
    
    doc.add_paragraph('Key Takeaways:')
    for takeaway in key_takeaways:
        doc.add_paragraph(takeaway, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('Next Steps:')
    next_steps = [
        'Start with the environment setup phase',
        'Create your first manual test case',
        'Convert it to an automated test',
        'Execute and view the test report',
        'Gradually build your test suite',
        'Integrate tests into your CI/CD pipeline',
    ]
    for step in next_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    # Save document
    output_path = r'd:\My Data\Automation-playwright\Playwright_Framework\Complete_Framework_Guide.docx'
    doc.save(output_path)
    print(f"Complete framework guide generated successfully: {output_path}")
    return output_path

if __name__ == "__main__":
    create_complete_report()
