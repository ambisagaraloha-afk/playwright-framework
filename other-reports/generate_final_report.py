from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def add_heading_with_style(doc, text, level, color_rgb=(0, 51, 102)):
    heading = doc.add_heading(text, level=level)
    heading.runs[0].font.color.rgb = RGBColor(*color_rgb)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_colored_paragraph(doc, text, color_rgb=(0, 0, 0), bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(*color_rgb)
    run.font.bold = bold
    run.font.size = Pt(size)
    return p

def shade_cell(cell, fill_color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_image_placeholder(doc, title, description):
    """Add an image placeholder section"""
    add_colored_paragraph(doc, f'[SCREENSHOT PLACEHOLDER]', bold=True, color_rgb=(255, 0, 0))
    doc.add_paragraph(f'Title: {title}')
    doc.add_paragraph(f'Description: {description}')
    doc.add_paragraph()

def create_final_comprehensive_report():
    doc = Document()
    
    # Setup margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # ===== TITLE PAGE =====
    title = doc.add_heading('Playwright Test Automation Framework', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph('From Manual Test Cases to Automated Scripts using Copilot Chat')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    info_para = doc.add_paragraph('Complete Methodology: Manual Testing → Excel → Copilot Chat → Automation Scripts')
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.runs[0].font.size = Pt(12)
    info_para.runs[0].italic = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Project Details
    details_table = doc.add_table(rows=6, cols=2)
    details_table.style = 'Light Grid Accent 1'
    
    details = [
        ('Document Type', 'Complete Framework Setup with Copilot Integration'),
        ('Audience', 'QA Engineers, Testers, Developers'),
        ('Framework', 'Playwright + TypeScript'),
        ('Key Tool', 'GitHub Copilot Chat with Instructions.md'),
        ('Created Date', datetime.now().strftime("%B %d, %Y")),
        ('Status', 'Production Ready'),
    ]
    
    for i, (key, value) in enumerate(details):
        key_cell = details_table.rows[i].cells[0]
        val_cell = details_table.rows[i].cells[1]
        key_cell.text = key
        val_cell.text = value
        shade_cell(key_cell, "4472C4")
        key_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    doc.add_page_break()
    
    # ===== TABLE OF CONTENTS =====
    add_heading_with_style(doc, 'Table of Contents', 1)
    toc_items = [
        '1. Introduction & Framework Overview',
        '2. System Requirements & Installation',
        '3. Phase 1: Environment Setup (with Screenshots)',
        '4. Phase 2: Manual Test Case Creation',
        '5. Phase 3: Creating Instructions.md for Copilot',
        '6. Phase 4: Using Copilot Chat to Automate Tests',
        '7. Phase 5: Project Structure & Organization',
        '8. Phase 6: Complete Automation Workflow',
        '9. Copilot Rules & Best Practices',
        '10. Step-by-Step Automation Example',
        '11. Testing & Execution',
        '12. Troubleshooting & Tips',
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 1. INTRODUCTION =====
    add_heading_with_style(doc, '1. Introduction & Framework Overview', 1)
    
    doc.add_paragraph(
        'This comprehensive guide demonstrates how to create end-to-end automated tests for web applications '
        'using Playwright and TypeScript, with AI assistance from GitHub Copilot Chat. The unique approach in this framework '
        'leverages an Instructions.md file to guide Copilot in generating high-quality, consistent, and maintainable test automation code.'
    )
    
    add_heading_with_style(doc, '1.1 The Three-Phase Automation Methodology', 2)
    
    phases_table = doc.add_table(rows=4, cols=2)
    phases_table.style = 'Light Grid Accent 1'
    
    header_cells = phases_table.rows[0].cells
    header_cells[0].text = 'Phase'
    header_cells[1].text = 'Description'
    shade_cell(header_cells[0], "4472C4")
    shade_cell(header_cells[1], "4472C4")
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    phases_data = [
        ('Phase 1: Manual Testing', 'Create detailed manual test cases in Excel with steps, data, and expected results'),
        ('Phase 2: Instructions Setup', 'Create Instructions.md file with Copilot rules, patterns, and best practices'),
        ('Phase 3: AI-Assisted Automation', 'Use Copilot Chat with Instructions.md to convert manual tests to automated scripts'),
    ]
    
    for i, (phase, desc) in enumerate(phases_data):
        row = phases_table.rows[i + 1]
        row.cells[0].text = phase
        row.cells[1].text = desc
    
    doc.add_page_break()
    
    # ===== 2. SYSTEM REQUIREMENTS & INSTALLATION =====
    add_heading_with_style(doc, '2. System Requirements & Installation', 1)
    
    add_heading_with_style(doc, '2.1 System Prerequisites', 2)
    
    req_table = doc.add_table(rows=5, cols=2)
    req_table.style = 'Light Grid Accent 1'
    
    header_cells = req_table.rows[0].cells
    header_cells[0].text = 'Requirement'
    header_cells[1].text = 'Specification'
    shade_cell(header_cells[0], "4472C4")
    shade_cell(header_cells[1], "4472C4")
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    reqs = [
        ('Operating System', 'Windows 10/11, macOS 11+, or Linux'),
        ('RAM', '8GB minimum (16GB recommended)'),
        ('Disk Space', '3GB free space'),
        ('Internet', 'Required for GitHub Copilot and npm packages'),
    ]
    
    for i, (req, spec) in enumerate(reqs):
        row = req_table.rows[i + 1]
        row.cells[0].text = req
        row.cells[1].text = spec
    
    add_heading_with_style(doc, '2.2 Required Software', 2)
    
    soft_table = doc.add_table(rows=8, cols=4)
    soft_table.style = 'Light Grid Accent 1'
    
    headers = ['Software', 'Version', 'Purpose', 'Download']
    for i, header_text in enumerate(soft_table.rows[0].cells):
        header_text.text = headers[i]
        shade_cell(header_text, "4472C4")
        header_text.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    software = [
        ('Node.js', '18.x LTS+', 'JavaScript runtime', 'nodejs.org'),
        ('VS Code', 'Latest', 'Code editor', 'code.visualstudio.com'),
        ('Playwright', '1.40+', 'Testing framework', 'npm install'),
        ('GitHub Copilot', 'Subscription', 'AI code assistant', 'github.com/copilot'),
        ('TypeScript', '5.x', 'Type-safe JavaScript', 'npm install'),
        ('Git', 'Latest', 'Version control', 'git-scm.com'),
        ('Excel Reader', '0.18.5', 'Read test cases', 'npm install xlsx'),
    ]
    
    for i, (soft, version, purpose, download) in enumerate(software):
        row = soft_table.rows[i + 1]
        row.cells[0].text = soft
        row.cells[1].text = version
        row.cells[2].text = purpose
        row.cells[3].text = download
    
    doc.add_page_break()
    
    # ===== 3. PHASE 1: ENVIRONMENT SETUP =====
    add_heading_with_style(doc, '3. Phase 1: Environment Setup (with Screenshots)', 1)
    
    add_heading_with_style(doc, '3.1 Step 1: Download and Install Node.js', 2)
    
    add_image_placeholder(doc, 'Node.js Official Website', 'Screenshot of nodejs.org homepage with LTS version button highlighted')
    
    node_steps = [
        'Go to https://nodejs.org/',
        'Click the LTS (Long Term Support) download button',
        'Run the installer and accept the license agreement',
        'Select default installation location',
        'Complete the installation',
    ]
    for step in node_steps:
        doc.add_paragraph(step, style='List Number')
    
    add_colored_paragraph(doc, 'Verify Installation:', bold=True, color_rgb=(0, 102, 204))
    doc.add_paragraph('node --version && npm --version')
    for run in doc.paragraphs[-1].runs:
        run.font.name = 'Courier New'
        run.font.bold = True
    
    add_image_placeholder(doc, 'Command Prompt Verification', 'Screenshot showing output: v18.17.0, npm 9.8.1')
    
    add_heading_with_style(doc, '3.2 Step 2: Install Visual Studio Code', 2)
    
    add_image_placeholder(doc, 'VS Code Download Page', 'Screenshot of code.visualstudio.com with download button')
    
    vscode_steps = [
        'Go to https://code.visualstudio.com/',
        'Click download for your operating system',
        'Run the installer',
        'Accept license and choose installation options',
        'Complete installation and launch VS Code',
    ]
    for step in vscode_steps:
        doc.add_paragraph(step, style='List Number')
    
    add_image_placeholder(doc, 'VS Code Welcome Screen', 'Screenshot of VS Code after first launch')
    
    add_heading_with_style(doc, '3.3 Step 3: Install Essential Extensions', 2)
    
    doc.add_paragraph('Open Extensions panel (Ctrl+Shift+X) and install:')
    
    extensions = [
        ('Playwright Test for VS Code', 'Run and debug tests directly in VS Code'),
        ('GitHub Copilot', 'AI-powered code suggestions'),
        ('GitHub Copilot Chat', 'Chat interface for AI assistance'),
        ('ESLint', 'Code quality checking'),
        ('Prettier', 'Code formatting'),
    ]
    
    for ext_name, purpose in extensions:
        doc.add_paragraph(f'{ext_name}: {purpose}', style='List Bullet')
    
    add_image_placeholder(doc, 'VS Code Extensions Panel', 'Screenshot showing Extensions panel with Playwright and Copilot extensions')
    
    add_heading_with_style(doc, '3.4 Step 4: Create Project Directory', 2)
    
    add_image_placeholder(doc, 'New Folder Creation', 'Screenshot of File Explorer creating new folder "Playwright_Framework"')
    
    add_colored_paragraph(doc, 'Steps:', bold=True, color_rgb=(0, 102, 204))
    steps_4 = [
        'Create folder: C:\\Users\\YourName\\Playwright_Framework',
        'Open VS Code',
        'File → Open Folder → Select the Playwright_Framework folder',
        'VS Code will open with the empty project',
    ]
    for step in steps_4:
        doc.add_paragraph(step, style='List Number')
    
    add_image_placeholder(doc, 'Project Opened in VS Code', 'Screenshot showing VS Code with empty Playwright_Framework folder')
    
    add_heading_with_style(doc, '3.5 Step 5: Open Terminal and Initialize Node Project', 2)
    
    add_image_placeholder(doc, 'Terminal in VS Code', 'Screenshot showing Terminal tab open at bottom of VS Code')
    
    add_colored_paragraph(doc, 'Initialize Node project:', bold=True, color_rgb=(0, 102, 204))
    cmd = 'npm init -y'
    doc.add_paragraph(cmd)
    for run in doc.paragraphs[-1].runs:
        run.font.name = 'Courier New'
        run.font.bold = True
    
    doc.add_paragraph('This creates package.json file')
    
    add_image_placeholder(doc, 'Package.json Created', 'Screenshot showing package.json file in project root')
    
    doc.add_page_break()
    
    # ===== 4. PHASE 2: MANUAL TEST CASE CREATION =====
    add_heading_with_style(doc, '4. Phase 2: Manual Test Case Creation', 1)
    
    doc.add_paragraph(
        'Manual test cases are the foundation of automation. They document the expected behavior of the application '
        'and serve as blueprints for automation scripts.'
    )
    
    add_heading_with_style(doc, '4.1 Test Case Template & Format', 2)
    
    add_heading_with_style(doc, 'Excel Columns Required:', 3)
    
    tc_template = [
        'TestCaseID - Unique identifier (TC-001)',
        'TestCaseName - Descriptive name of the test',
        'Module - Feature/Module being tested',
        'Description - Detailed description',
        'Priority - High/Medium/Low',
        'Severity - Critical/Major/Minor',
        'Preconditions - Setup requirements',
        'TestData - Input values and test data',
        'Steps - Numbered execution steps',
        'ExpectedResult - Expected outcome',
        'PostConditions - Cleanup steps',
        'Status - Pass/Fail/Not Executed',
    ]
    
    for item in tc_template:
        doc.add_paragraph(item, style='List Bullet')
    
    add_image_placeholder(doc, 'Excel Test Case Sheet', 'Screenshot of Excel file with test case columns and sample data rows')
    
    add_heading_with_style(doc, '4.2 Example Manual Test Cases', 2)
    
    add_colored_paragraph(doc, 'Example 1: User Login Test', bold=True, color_rgb=(0, 102, 204))
    
    tc1_table = doc.add_table(rows=10, cols=2)
    tc1_table.style = 'Light Grid'
    
    tc1_data = [
        ('TestCaseID', 'TC-001'),
        ('TestCaseName', 'Verify valid user login'),
        ('Module', 'Authentication'),
        ('Priority', 'High'),
        ('Preconditions', 'User account exists; User not logged in'),
        ('TestData', 'Username: testuser@example.com; Password: Test@123'),
        ('Steps', '1. Navigate to login page\n2. Enter email\n3. Enter password\n4. Click Login'),
        ('ExpectedResult', 'User logs in successfully; Dashboard displays'),
        ('Status', 'Not Executed'),
    ]
    
    for i, (key, value) in enumerate(tc1_data):
        row = tc1_table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
        shade_cell(row.cells[0], "D9E1F2")
    
    add_image_placeholder(doc, 'Excel Test Case Entry', 'Screenshot of TC-001 entry in Excel spreadsheet')
    
    doc.add_page_break()
    
    # ===== 5. PHASE 3: CREATING INSTRUCTIONS.MD FOR COPILOT =====
    add_heading_with_style(doc, '5. Phase 3: Creating Instructions.md for Copilot', 1)
    
    doc.add_paragraph(
        'The Instructions.md file is crucial - it tells GitHub Copilot Chat how to generate automation scripts. '
        'It contains rules, patterns, best practices, and DemoBlaze application details.'
    )
    
    add_heading_with_style(doc, '5.1 Why Instructions.md Matters', 2)
    
    doc.add_paragraph(
        'GitHub Copilot is an AI model that learns from patterns. By providing clear instructions, rules, and examples, '
        'Copilot generates consistent, high-quality code that follows your framework\'s conventions.'
    )
    
    benefits = [
        'Ensures consistent code quality and style',
        'Guides Copilot to use Page Object Model correctly',
        'Prevents common mistakes in test automation',
        'Reduces manual code review and fixes',
        'Speeds up automation script generation',
        'Maintains framework standards across all tests',
    ]
    for benefit in benefits:
        doc.add_paragraph(benefit, style='List Bullet')
    
    add_heading_with_style(doc, '5.2 Instructions.md File Structure', 2)
    
    add_colored_paragraph(doc, 'Create file: Instructions_DemoBlaze.md at project root', bold=True, color_rgb=(0, 102, 204))
    
    add_image_placeholder(doc, 'Instructions File Created', 'Screenshot showing Instructions_DemoBlaze.md file in VS Code project root')
    
    add_heading_with_style(doc, '5.3 Key Sections in Instructions.md', 2)
    
    sections = {
        'Quick Start': 'Explains how user provides test case info',
        'Target Application Details': 'URL, features, and capabilities of the app being tested',
        'Test Case Source Format': 'Expected Excel format and columns',
        'Critical Rules': 'Mandatory rules for Copilot to follow',
        'Rule Examples': 'Before/After code examples showing correct patterns',
        'Application-Specific Selectors': 'DemoBlaze page locators and element IDs',
        'Test Naming Conventions': 'How to name test methods and classes',
    }
    
    for section, purpose in sections.items():
        doc.add_paragraph(f'{section}: {purpose}', style='List Bullet')
    
    add_heading_with_style(doc, '5.4 Critical Rules to Include in Instructions.md', 2)
    
    critical_rules = [
        'Rule 1: Always Display Test Case First AND Wait for Confirmation',
        'Rule 2: Capture BOTH Action AND Validation Locators',
        'Rule 3: Never Guess Locators — Verify First',
        'Rule 4: Use Locator Priority (data-testid > id > class > XPath)',
        'Rule 5: Use Page Object Model with Proper Structure',
        'Rule 6: Element Names Must Use "Locator" Suffix',
        'Rule 7: Avoid Complex Assertions in Test Methods',
        'Rule 8: No Hard Waits — Use Playwright Built-in Waits',
        'Rule 9: Implement ALL Test Steps — No Partial Implementation',
        'Rule 10: Follow Test Method Naming Convention (e.g., demoBlaze_login_verifyValidUserLogin)',
    ]
    
    for rule in critical_rules:
        doc.add_paragraph(rule, style='List Bullet')
    
    add_image_placeholder(doc, 'Instructions.md Content', 'Screenshot showing Instructions_DemoBlaze.md file open in VS Code with all critical rules visible')
    
    doc.add_page_break()
    
    # ===== 6. PHASE 4: USING COPILOT CHAT =====
    add_heading_with_style(doc, '6. Phase 4: Using Copilot Chat to Automate Tests', 1)
    
    doc.add_paragraph(
        'This is where the magic happens. GitHub Copilot Chat uses the Instructions.md file to generate '
        'high-quality, framework-compliant test automation code.'
    )
    
    add_heading_with_style(doc, '6.1 Step 1: Upload Instructions.md to Copilot Chat', 2)
    
    doc.add_paragraph('First, provide the Instructions.md context to Copilot:')
    
    steps_upload = [
        'Open Copilot Chat in VS Code (Ctrl+Shift+I or View → Copilot Chat)',
        'Drag and drop Instructions_DemoBlaze.md into the chat',
        'Or use the @ symbol: @Instructions_DemoBlaze.md',
        'Copilot will read and understand the instructions',
    ]
    
    for step in steps_upload:
        doc.add_paragraph(step, style='List Number')
    
    add_image_placeholder(doc, 'Copilot Chat Open', 'Screenshot of Copilot Chat panel opened with Instructions.md being dragged in')
    
    add_heading_with_style(doc, '6.2 Step 2: Provide Test Case Details to Copilot', 2)
    
    doc.add_paragraph('Format your request to Copilot as:')
    
    sample_prompt = '''Using the Instructions_DemoBlaze.md file, please automate TC-001:
- TestCaseID: TC-001
- TestCaseName: Verify valid user login
- Module: Authentication
- Preconditions: User account exists, User not logged in
- TestData: Username: testuser@example.com, Password: Test@123
- Steps:
  1. Navigate to login page
  2. Enter email
  3. Enter password
  4. Click Login button
- ExpectedResult: User logs in successfully, Dashboard displays

Please provide the complete Page Object and Test Specification.'''
    
    doc.add_paragraph(sample_prompt, style='No Spacing')
    for run in doc.paragraphs[-1].runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    add_image_placeholder(doc, 'Copilot Chat Prompt', 'Screenshot showing test case details entered in Copilot Chat')
    
    add_heading_with_style(doc, '6.3 Step 3: Copilot Generates the Code', 2)
    
    doc.add_paragraph('Copilot will generate:')
    
    generated_items = [
        'Login Page Object (loginPage.ts) with locators and methods',
        'Test Specification (auth.spec.ts) with test cases',
        'Complete TypeScript code following all Instructions.md rules',
    ]
    
    for item in generated_items:
        doc.add_paragraph(item, style='List Bullet')
    
    add_image_placeholder(doc, 'Copilot Generated Code', 'Screenshot showing Copilot Chat with generated Page Object code')
    
    add_heading_with_style(doc, '6.4 Step 4: Review and Adjust', 2)
    
    review_steps = [
        'Review the generated code for accuracy',
        'Verify locators are correct for DemoBlaze application',
        'Check if all test steps are implemented',
        'Make any necessary adjustments or ask Copilot to fix issues',
    ]
    
    for step in review_steps:
        doc.add_paragraph(step, style='List Number')
    
    add_heading_with_style(doc, '6.5 Step 5: Create Files and Copy Code', 2)
    
    doc.add_paragraph('Create the following files in your project:')
    
    file_steps = [
        'Pages/loginPage.ts → Paste the Page Object code',
        'tests/auth.spec.ts → Paste the Test Specification code',
        'utilities/testData.ts → Add test credentials and data',
    ]
    
    for step in file_steps:
        doc.add_paragraph(step, style='List Number')
    
    add_image_placeholder(doc, 'Files Created in VS Code', 'Screenshot showing Pages, tests, and utilities folders with new files')
    
    doc.add_page_break()
    
    # ===== 7. PHASE 5: PROJECT STRUCTURE =====
    add_heading_with_style(doc, '7. Phase 5: Project Structure & Organization', 1)
    
    add_heading_with_style(doc, '7.1 Complete Folder Structure', 2)
    
    structure_text = '''Playwright_Framework/
├── Pages/
│   ├── basePage.ts
│   ├── loginPage.ts
│   ├── homePage.ts
│   └── cartPage.ts
│
├── tests/
│   ├── auth.spec.ts
│   ├── shopping.spec.ts
│   ├── checkout.spec.ts
│   └── regression.spec.ts
│
├── utilities/
│   ├── testData.ts
│   ├── helpers.ts
│   └── logger.ts
│
├── test-cases_manual/
│   ├── TestCases.xlsx
│   └── TestCases_Completed.xlsx
│
├── Instructions_DemoBlaze.md
├── playwright.config.ts
├── package.json
├── tsconfig.json
└── README.md'''
    
    doc.add_paragraph(structure_text, style='No Spacing')
    for run in doc.paragraphs[-1].runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    add_image_placeholder(doc, 'Project Structure in Explorer', 'Screenshot of VS Code Explorer showing complete folder structure')
    
    add_heading_with_style(doc, '7.2 Folder Responsibilities', 2)
    
    folders = {
        'Pages/': 'Page Object Model classes containing locators and methods for each page',
        'tests/': 'Test specification files containing test cases organized by feature',
        'utilities/': 'Helper functions, test data constants, and logging utilities',
        'test-cases_manual/': 'Excel files with manual test cases before and after automation',
    }
    
    for folder, responsibility in folders.items():
        doc.add_paragraph(f'{folder}: {responsibility}', style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 8. PHASE 6: COMPLETE AUTOMATION WORKFLOW =====
    add_heading_with_style(doc, '8. Phase 6: Complete Automation Workflow', 1)
    
    add_heading_with_style(doc, '8.1 From Manual Test to Automated Test', 2)
    
    workflow_steps = [
        ('Manual Test Case Created', 'Document test in Excel with all details (steps, data, expected results)'),
        ('Review & Approve', 'Test stakeholders review and approve the manual test case'),
        ('Copilot Request', 'Use Copilot Chat with Instructions.md to generate automation'),
        ('Code Generated', 'Copilot generates Page Object and Test Specification'),
        ('Code Review', 'Review generated code, verify locators, check implementation'),
        ('Execute Test', 'Run the automated test to verify it works correctly'),
        ('Update TC Status', 'Mark test case as "Automated" in Excel'),
        ('Maintenance', 'Keep tests updated with application changes'),
    ]
    
    for i, (step_name, step_desc) in enumerate(workflow_steps):
        doc.add_paragraph(f'Step {i+1}: {step_name}', style='List Number')
        doc.add_paragraph(f'{step_desc}', style='List Bullet')
    
    add_image_placeholder(doc, 'Workflow Diagram', 'Screenshot or diagram showing flow from Manual TC → Excel → Copilot Chat → Automation → Test Reports')
    
    doc.add_page_break()
    
    # ===== 9. COPILOT RULES & BEST PRACTICES =====
    add_heading_with_style(doc, '9. Copilot Rules & Best Practices', 1)
    
    doc.add_paragraph(
        'These are the critical rules that must be included in your Instructions.md file. '
        'They ensure Copilot generates consistent, maintainable code.'
    )
    
    add_heading_with_style(doc, '9.1 The 10 Critical Rules', 2)
    
    rules = {
        'Rule 1: Always Display Test Case First': [
            'Display complete test case in chat',
            'WAIT for user to reply "Yes, proceed"',
            'DO NOT start coding without confirmation',
        ],
        'Rule 2: Capture Both Action & Validation Locators': [
            'Capture action element locators (inputs, buttons)',
            'Capture validation locators (messages, headers, badges)',
            'Document ALL locators before coding',
        ],
        'Rule 3: Never Guess Locators': [
            'Inspect actual DOM in MCP browser',
            'Verify selector works before using',
            'Use unique, reliable selectors',
        ],
        'Rule 4: Use Locator Priority': [
            '1. data-testid attribute (preferred)',
            '2. id attribute',
            '3. class attribute',
            '4. CSS selectors',
            '5. XPath (last resort)',
        ],
        'Rule 5: Use Page Object Model': [
            'Declare page elements as class properties',
            'Use methods for page interactions',
            'Separate test logic from UI locators',
        ],
        'Rule 6: Element Names Must Use Locator Suffix': [
            'YES: loginButtonLocator, usernameFieldLocator',
            'NO: loginButton, usernameField',
        ],
        'Rule 7: Avoid Complex Assertions': [
            'Create methods in page objects for assertions',
            'Keep test files focused on test logic',
            'Return boolean values from page methods',
        ],
        'Rule 8: No Hard Waits': [
            'Use waitFor() instead of setTimeout()',
            'Wait for element visibility',
            'Set appropriate timeouts',
        ],
        'Rule 9: Implement ALL Steps': [
            'Complete all test steps',
            'Implement all validations',
            'No partial implementations or TODOs',
        ],
        'Rule 10: Follow Naming Convention': [
            'Pattern: demoBlaze_[page]_[description]()',
            'Example: demoBlaze_login_verifyValidUserLogin()',
            'Example: demoBlaze_cart_addSingleProduct()',
        ],
    }
    
    for rule_title, rule_details in rules.items():
        add_colored_paragraph(doc, rule_title, bold=True, color_rgb=(0, 102, 204))
        for detail in rule_details:
            doc.add_paragraph(detail, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 10. STEP-BY-STEP AUTOMATION EXAMPLE =====
    add_heading_with_style(doc, '10. Step-by-Step Automation Example', 1)
    
    doc.add_paragraph(
        'This section shows a complete real-world example from manual test case to running automated tests.'
    )
    
    add_heading_with_style(doc, '10.1 Example: Add to Cart Test', 2)
    
    add_colored_paragraph(doc, 'Step 1: Create Manual Test Case in Excel', bold=True, color_rgb=(0, 102, 204))
    
    manual_tc = '''
TestCaseID: TC-002
TestCaseName: Verify user can add single product to cart
Module: Shopping Cart
Priority: High
Preconditions: User is on DemoBlaze home page
TestData: 
  - Product: Samsung Galaxy S10
  - Quantity: 1
Steps:
  1. Navigate to https://demoblaze.com/index.html
  2. Look for Samsung Galaxy S10 product
  3. Click "Add to cart" button
  4. Verify success message appears
  5. Verify product is added to cart (check cart count)
ExpectedResult:
  - Product is successfully added to cart
  - Cart count increases by 1
  - Success message is displayed
Status: Not Executed
    '''
    
    doc.add_paragraph(manual_tc, style='No Spacing')
    for run in doc.paragraphs[-1].runs:
        run.font.size = Pt(10)
    
    add_image_placeholder(doc, 'Manual TC in Excel', 'Screenshot of TC-002 added to Excel spreadsheet with all columns filled')
    
    add_colored_paragraph(doc, 'Step 2: Request Copilot to Automate TC-002', bold=True, color_rgb=(0, 102, 204))
    
    copilot_request = '''Using Instructions_DemoBlaze.md, please automate TC-002:

TestCaseID: TC-002
TestCaseName: Verify user can add single product to cart
Module: Shopping Cart
Priority: High
Preconditions: User is on DemoBlaze home page
TestData: Product: Samsung Galaxy S10, Quantity: 1
Steps:
  1. Navigate to https://demoblaze.com/index.html
  2. Look for Samsung Galaxy S10 product
  3. Click "Add to cart" button
  4. Verify success message appears
  5. Verify product is added to cart

ExpectedResult: Product successfully added, Cart count increased, Success message displayed

Please generate the homePage.ts Page Object class and the shopping.spec.ts test file.'''
    
    doc.add_paragraph(copilot_request, style='No Spacing')
    for run in doc.paragraphs[-1].runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    add_image_placeholder(doc, 'Copilot Chat Request', 'Screenshot of the automation request pasted in Copilot Chat')
    
    add_colored_paragraph(doc, 'Step 3: Copilot Generates Code', bold=True, color_rgb=(0, 102, 204))
    
    generated_code = '''// Pages/homePage.ts
import { Page, Locator } from '@playwright/test';

export class HomePage {
    readonly page: Page;
    readonly samsungGalaxyS10Locator: Locator;
    readonly addToCartButtonLocator: Locator;
    readonly successMessageLocator: Locator;
    readonly cartBadgeLocator: Locator;

    constructor(page: Page) {
        this.page = page;
        this.samsungGalaxyS10Locator = page.locator('a:has-text("Samsung galaxy s6")');
        this.addToCartButtonLocator = page.locator('a:has-text("Add to cart")');
        this.successMessageLocator = page.locator('.alert-success');
        this.cartBadgeLocator = page.locator('#cartur');
    }

    async addProductToCart(productName: string) {
        await this.samsungGalaxyS10Locator.click();
        await this.page.waitForTimeout(500);
        await this.addToCartButtonLocator.click();
    }

    async isSuccessMessageDisplayed(): Promise<boolean> {
        return await this.successMessageLocator.isVisible();
    }
}

// tests/shopping.spec.ts
import { test, expect } from '@playwright/test';
import { HomePage } from '../Pages/homePage';

test.describe('Shopping Cart Tests', () => {
    let homePage: HomePage;

    test.beforeEach(async ({ page }) => {
        await page.goto('https://demoblaze.com/index.html');
        homePage = new HomePage(page);
    });

    test('demoBlaze_cart_addSingleProduct', async ({ page }) => {
        // Add product to cart
        await homePage.addProductToCart('Samsung Galaxy S10');
        
        // Verify success message
        const isSuccessDisplayed = await homePage.isSuccessMessageDisplayed();
        expect(isSuccessDisplayed).toBe(true);
    });
});'''
    
    doc.add_paragraph(generated_code, style='No Spacing')
    for run in doc.paragraphs[-1].runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
    
    add_image_placeholder(doc, 'Copilot Generated Code', 'Screenshot of Copilot Chat showing generated Page Object and Test files')
    
    add_colored_paragraph(doc, 'Step 4: Create Files and Copy Code', bold=True, color_rgb=(0, 102, 204))
    
    create_steps = [
        'Create Pages/homePage.ts with Page Object code',
        'Create tests/shopping.spec.ts with test specifications',
        'Save both files',
    ]
    
    for step in create_steps:
        doc.add_paragraph(step, style='List Number')
    
    add_image_placeholder(doc, 'Files Created', 'Screenshot showing homePage.ts and shopping.spec.ts created in project')
    
    add_colored_paragraph(doc, 'Step 5: Execute the Test', bold=True, color_rgb=(0, 102, 204))
    
    execute_cmd = 'npx playwright test tests/shopping.spec.ts --headed'
    doc.add_paragraph(execute_cmd)
    for run in doc.paragraphs[-1].runs:
        run.font.name = 'Courier New'
        run.font.bold = True
    
    add_image_placeholder(doc, 'Test Execution', 'Screenshot of terminal showing test execution with browser visible running the test')
    
    add_colored_paragraph(doc, 'Step 6: View Test Report', bold=True, color_rgb=(0, 102, 204))
    
    report_cmd = 'npx playwright show-report'
    doc.add_paragraph(report_cmd)
    for run in doc.paragraphs[-1].runs:
        run.font.name = 'Courier New'
        run.font.bold = True
    
    add_image_placeholder(doc, 'Test Report', 'Screenshot of HTML report showing TC-002 passed with execution details')
    
    doc.add_page_break()
    
    # ===== 11. TESTING & EXECUTION =====
    add_heading_with_style(doc, '11. Testing & Execution', 1)
    
    add_heading_with_style(doc, '11.1 Running Tests', 2)
    
    test_commands = [
        ('All Tests', 'npx playwright test'),
        ('Specific File', 'npx playwright test tests/auth.spec.ts'),
        ('Specific Test', 'npx playwright test -g "verify valid user"'),
        ('Headed Mode', 'npx playwright test --headed'),
        ('Debug Mode', 'npx playwright test --debug'),
        ('Chromium Only', 'npx playwright test --project=chromium'),
        ('Verbose Output', 'npx playwright test --verbose'),
    ]
    
    for cmd_name, cmd in test_commands:
        doc.add_paragraph(f'{cmd_name}:')
        cmd_para = doc.add_paragraph(cmd)
        for run in cmd_para.runs:
            run.font.name = 'Courier New'
            run.font.bold = True
    
    add_image_placeholder(doc, 'Terminal Test Execution', 'Screenshot showing tests running in terminal with progress output')
    
    add_heading_with_style(doc, '11.2 Viewing Test Reports', 2)
    
    doc.add_paragraph('View the HTML report:')
    report_para = doc.add_paragraph('npx playwright show-report')
    for run in report_para.runs:
        run.font.name = 'Courier New'
        run.font.bold = True
    
    doc.add_paragraph('The report shows:')
    report_features = [
        'Test results summary (passed, failed, skipped)',
        'Detailed test execution logs',
        'Screenshots of test execution',
        'Videos of failed tests',
        'Trace files for debugging',
        'Timing and performance metrics',
    ]
    for feature in report_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    add_image_placeholder(doc, 'HTML Test Report', 'Screenshot of HTML report page showing test results, screenshots, and videos')
    
    doc.add_page_break()
    
    # ===== 12. TROUBLESHOOTING & TIPS =====
    add_heading_with_style(doc, '12. Troubleshooting & Tips', 1)
    
    add_heading_with_style(doc, '12.1 Common Issues and Solutions', 2)
    
    issues_table = doc.add_table(rows=8, cols=2)
    issues_table.style = 'Light Grid Accent 1'
    
    header_cells = issues_table.rows[0].cells
    header_cells[0].text = 'Issue'
    header_cells[1].text = 'Solution'
    shade_cell(header_cells[0], "4472C4")
    shade_cell(header_cells[1], "4472C4")
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    issues = [
        ('Element not found', 'Check locator is correct, use page.screenshot() to debug'),
        ('Test timeout', 'Increase timeout in playwright.config.ts, add proper waits'),
        ('Flaky tests', 'Use waitFor() with state conditions, avoid hard waits'),
        ('Copilot generates incorrect code', 'Provide clear Instructions.md, ask Copilot to refactor'),
        ('Module not found error', 'Run npm install, check import paths'),
        ('Tests run too slow', 'Enable parallelization in config, optimize waits'),
        ('Report not generating', 'Check HTML reporter in config, verify permissions'),
    ]
    
    for i, (issue, solution) in enumerate(issues):
        row = issues_table.rows[i + 1]
        row.cells[0].text = issue
        row.cells[1].text = solution
    
    add_heading_with_style(doc, '12.2 Tips for Success', 2)
    
    tips = [
        'Keep Instructions.md updated with new patterns and rules',
        'Review Copilot-generated code carefully before using',
        'Use meaningful variable names (ending with "Locator")',
        'Test generated code in headed mode first',
        'Maintain test data separately in utilities/testData.ts',
        'Use Page Object Model consistently across all tests',
        'Add comments to complex test logic',
        'Run tests regularly to catch failures early',
        'Update locators when UI changes occur',
        'Use version control (Git) for all test code',
    ]
    
    for tip in tips:
        doc.add_paragraph(tip, style='List Bullet')
    
    add_heading_with_style(doc, '12.3 Quick Command Reference', 2)
    
    quick_ref = [
        'npm init -y - Initialize new Node project',
        'npm install --save-dev @playwright/test - Install Playwright',
        'npx playwright install - Download browser binaries',
        'npx playwright test - Run all tests',
        'npx playwright test --headed - Run with visible browser',
        'npx playwright test --debug - Debug mode with Inspector',
        'npx playwright show-report - View test report',
        'npx playwright codegen - Record UI interactions',
    ]
    
    for ref in quick_ref:
        doc.add_paragraph(ref, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== APPENDIX =====
    add_heading_with_style(doc, 'Appendix: Sample Instructions.md Content', 1)
    
    doc.add_paragraph('Here is an example of the critical rules section from Instructions.md:')
    
    sample_instructions = '''## 🚨 CRITICAL RULES — MANDATORY

### RULE #1: Always Display Test Case First AND Wait for Confirmation
❌ **DON'T:** Read Excel and start coding immediately
✅ **DO:** Display complete test case in chat → WAIT for user reply

Process:
1. Read Excel test case
2. Display ALL test details in chat
3. Say: "I've read the test case. Please confirm if this is correct."
4. **STOP** — Do not code until confirmation
5. Only after confirmation → Continue

### RULE #2: Capture BOTH Action AND Validation Locators
❌ **DON'T:** Only capture input/button locators
✅ **DO:** Also capture error messages, success indicators, page headers

### RULE #5: Use Page Object Model with Proper Structure
❌ **DON'T:**
await page.locator('[some-selector]').click();

✅ **DO:**
readonly addToCartButtonLocator = page.locator('[data-id="add-to-cart"]');
async addProductToCart() {
  await this.addToCartButtonLocator.click();
}

### RULE #10: Test Method Naming Convention
❌ **DON'T:** test1(), loginTest(), TC001()
✅ **DO:** demoBlaze_login_verifyValidUserLogin()
Pattern: demoBlaze_[page]_[testDescription]()'''
    
    doc.add_paragraph(sample_instructions, style='No Spacing')
    for run in doc.paragraphs[-1].runs:
        run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # ===== CONCLUSION =====
    add_heading_with_style(doc, 'Conclusion', 1)
    
    doc.add_paragraph(
        'You now have a complete, comprehensive guide for building and maintaining a Playwright test automation framework '
        'with GitHub Copilot Chat assistance. By following this methodology, you can:'
    )
    
    conclusion_points = [
        'Create well-structured manual test cases in Excel',
        'Use Instructions.md to guide Copilot in code generation',
        'Generate consistent, maintainable automation scripts',
        'Execute tests and view comprehensive HTML reports',
        'Scale your test suite efficiently',
        'Maintain code quality and framework standards',
    ]
    
    for point in conclusion_points:
        doc.add_paragraph(point, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('Next Steps:')
    next_steps = [
        'Set up your environment following Phase 1-3',
        'Create your Instructions_DemoBlaze.md file',
        'Write your first manual test case in Excel',
        'Use Copilot Chat to generate automation code',
        'Execute and verify the test',
        'Gradually build your test suite',
        'Integrate into CI/CD pipeline',
    ]
    
    for step in next_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    # Save document
    output_path = r'd:\My Data\Automation-playwright\Playwright_Framework\FINAL_Complete_Framework_Guide.docx'
    doc.save(output_path)
    print(f"Final comprehensive guide generated: {output_path}")
    return output_path

if __name__ == "__main__":
    create_final_comprehensive_report()
