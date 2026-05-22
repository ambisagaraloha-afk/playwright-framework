from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from PIL import Image, ImageDraw, ImageFont
import io

def create_placeholder_image(title, description, width=800, height=400):
    """Create a simple placeholder image with title and description"""
    img = Image.new('RGB', (width, height), color=(240, 248, 255))
    draw = ImageDraw.Draw(img)
    
    # Try to use a built-in font, fallback to default
    try:
        title_font = ImageFont.truetype("arial.ttf", 24)
        text_font = ImageFont.truetype("arial.ttf", 14)
    except:
        title_font = ImageFont.load_default()
        text_font = ImageFont.load_default()
    
    # Draw border
    draw.rectangle([10, 10, width-10, height-10], outline=(0, 102, 204), width=3)
    
    # Draw title
    draw.text((30, 30), f"[SCREENSHOT: {title}]", fill=(0, 51, 102), font=title_font)
    
    # Draw description
    draw.text((30, 100), description, fill=(0, 0, 0), font=text_font)
    
    # Draw placeholder text
    draw.text((30, 200), "This is a placeholder for:", fill=(100, 100, 100), font=text_font)
    draw.text((30, 240), "• Insert your actual screenshot here", fill=(100, 100, 100), font=text_font)
    draw.text((30, 280), "• Or use this diagram as reference", fill=(100, 100, 100), font=text_font)
    
    # Save to bytes
    img_bytes = io.BytesIO()
    img.save(img_bytes, format='PNG')
    img_bytes.seek(0)
    return img_bytes

def add_heading_with_style(doc, text, level, color_rgb=(0, 51, 102)):
    heading = doc.add_heading(text, level=level)
    heading.runs[0].font.color.rgb = RGBColor(*color_rgb)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_colored_paragraph(doc, text, color_rgb=(0, 0, 0), bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(*color_rgb)
    run.font.bold = bold
    run.font.size = Pt(size)
    return p

def shade_cell(cell, fill_color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_screenshot_section(doc, title, description, img_width=6.5):
    """Add a screenshot placeholder section"""
    add_colored_paragraph(doc, title, bold=True, color_rgb=(0, 102, 204), size=12)
    doc.add_paragraph(description)
    
    # Add placeholder image
    img_bytes = create_placeholder_image(title, description)
    doc.add_picture(img_bytes, width=Inches(img_width))
    
    # Center the image
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing

def create_final_document_with_images():
    doc = Document()
    
    # Setup margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # ===== TITLE PAGE =====
    title = doc.add_heading('Playwright Test Automation Framework', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph('From Manual Test Cases to Automated Scripts using Copilot Chat')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    info_para = doc.add_paragraph('Complete Methodology: Manual Testing → Excel → Copilot Chat → Automation Scripts')
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.runs[0].font.size = Pt(12)
    info_para.runs[0].italic = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Project Details
    details_table = doc.add_table(rows=6, cols=2)
    details_table.style = 'Light Grid Accent 1'
    
    details = [
        ('Document Type', 'Complete Framework Setup with Copilot Integration'),
        ('Audience', 'QA Engineers, Testers, Developers'),
        ('Framework', 'Playwright + TypeScript'),
        ('Key Tool', 'GitHub Copilot Chat with Instructions.md'),
        ('Created Date', datetime.now().strftime("%B %d, %Y")),
        ('Status', 'Production Ready with Screenshots'),
    ]
    
    for i, (key, value) in enumerate(details):
        key_cell = details_table.rows[i].cells[0]
        val_cell = details_table.rows[i].cells[1]
        key_cell.text = key
        val_cell.text = value
        shade_cell(key_cell, "4472C4")
        key_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    doc.add_page_break()
    
    # ===== TABLE OF CONTENTS =====
    add_heading_with_style(doc, 'Table of Contents', 1)
    toc_items = [
        '1. Introduction & Framework Overview',
        '2. System Requirements & Installation',
        '3. Phase 1: Environment Setup',
        '4. Phase 2: Manual Test Case Creation',
        '5. Phase 3: Creating Instructions.md for Copilot',
        '6. Phase 4: Using Copilot Chat to Automate Tests',
        '7. Phase 5: Project Structure & Organization',
        '8. Phase 6: Complete Automation Workflow',
        '9. Copilot Rules & Best Practices',
        '10. Step-by-Step Automation Example',
        '11. Testing & Execution',
        '12. Troubleshooting & Tips',
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 1. INTRODUCTION =====
    add_heading_with_style(doc, '1. Introduction & Framework Overview', 1)
    
    doc.add_paragraph(
        'This comprehensive guide demonstrates how to create end-to-end automated tests for web applications '
        'using Playwright and TypeScript, with AI assistance from GitHub Copilot Chat. The unique approach in this framework '
        'leverages an Instructions.md file to guide Copilot in generating high-quality, consistent, and maintainable test automation code.'
    )
    
    add_heading_with_style(doc, '1.1 The Three-Phase Automation Methodology', 2)
    
    phases_table = doc.add_table(rows=4, cols=2)
    phases_table.style = 'Light Grid Accent 1'
    
    header_cells = phases_table.rows[0].cells
    header_cells[0].text = 'Phase'
    header_cells[1].text = 'Description'
    shade_cell(header_cells[0], "4472C4")
    shade_cell(header_cells[1], "4472C4")
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    phases_data = [
        ('Phase 1: Manual Testing', 'Create detailed manual test cases in Excel with steps, data, and expected results'),
        ('Phase 2: Instructions Setup', 'Create Instructions.md file with Copilot rules, patterns, and best practices'),
        ('Phase 3: AI-Assisted Automation', 'Use Copilot Chat with Instructions.md to convert manual tests to automated scripts'),
    ]
    
    for i, (phase, desc) in enumerate(phases_data):
        row = phases_table.rows[i + 1]
        row.cells[0].text = phase
        row.cells[1].text = desc
    
    doc.add_page_break()
    
    # ===== 2. SYSTEM REQUIREMENTS =====
    add_heading_with_style(doc, '2. System Requirements & Installation', 1)
    
    add_heading_with_style(doc, '2.1 System Prerequisites', 2)
    
    req_table = doc.add_table(rows=5, cols=2)
    req_table.style = 'Light Grid Accent 1'
    
    header_cells = req_table.rows[0].cells
    header_cells[0].text = 'Requirement'
    header_cells[1].text = 'Specification'
    shade_cell(header_cells[0], "4472C4")
    shade_cell(header_cells[1], "4472C4")
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    reqs = [
        ('Operating System', 'Windows 10/11, macOS 11+, or Linux'),
        ('RAM', '8GB minimum (16GB recommended)'),
        ('Disk Space', '3GB free space'),
        ('Internet', 'Required for GitHub Copilot and npm packages'),
    ]
    
    for i, (req, spec) in enumerate(reqs):
        row = req_table.rows[i + 1]
        row.cells[0].text = req
        row.cells[1].text = spec
    
    add_heading_with_style(doc, '2.2 Required Software', 2)
    
    soft_table = doc.add_table(rows=8, cols=4)
    soft_table.style = 'Light Grid Accent 1'
    
    headers = ['Software', 'Version', 'Purpose', 'Download']
    for i, header_text in enumerate(soft_table.rows[0].cells):
        header_text.text = headers[i]
        shade_cell(header_text, "4472C4")
        header_text.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    software = [
        ('Node.js', '18.x LTS+', 'JavaScript runtime', 'nodejs.org'),
        ('VS Code', 'Latest', 'Code editor', 'code.visualstudio.com'),
        ('Playwright', '1.40+', 'Testing framework', 'npm install'),
        ('GitHub Copilot', 'Subscription', 'AI code assistant', 'github.com/copilot'),
        ('TypeScript', '5.x', 'Type-safe JavaScript', 'npm install'),
        ('Git', 'Latest', 'Version control', 'git-scm.com'),
        ('Excel Reader', '0.18.5', 'Read test cases', 'npm install xlsx'),
    ]
    
    for i, (soft, version, purpose, download) in enumerate(software):
        row = soft_table.rows[i + 1]
        row.cells[0].text = soft
        row.cells[1].text = version
        row.cells[2].text = purpose
        row.cells[3].text = download
    
    doc.add_page_break()
    
    # ===== 3. PHASE 1: ENVIRONMENT SETUP =====
    add_heading_with_style(doc, '3. Phase 1: Environment Setup with Screenshots', 1)
    
    add_heading_with_style(doc, '3.1 Step 1: Download and Install Node.js', 2)
    add_screenshot_section(doc, 'Node.js Official Website', 
        'Screenshot showing nodejs.org homepage with LTS version download button highlighted')
    
    node_steps = [
        'Go to https://nodejs.org/',
        'Click the LTS (Long Term Support) download button',
        'Run the installer and accept the license agreement',
        'Select default installation location',
        'Complete the installation',
    ]
    for step in node_steps:
        doc.add_paragraph(step, style='List Number')
    
    add_colored_paragraph(doc, 'Verify Installation:', bold=True, color_rgb=(0, 102, 204))
    doc.add_paragraph('node --version && npm --version')
    for run in doc.paragraphs[-1].runs:
        run.font.name = 'Courier New'
        run.font.bold = True
    
    add_screenshot_section(doc, 'Command Prompt Verification', 
        'Screenshot showing terminal output: v18.17.0, npm 9.8.1 etc.')
    
    add_heading_with_style(doc, '3.2 Step 2: Install Visual Studio Code', 2)
    add_screenshot_section(doc, 'VS Code Download Page', 
        'Screenshot of code.visualstudio.com homepage with download button visible')
    
    vscode_steps = [
        'Go to https://code.visualstudio.com/',
        'Click download for your operating system',
        'Run the installer',
        'Accept license and choose installation options',
        'Complete installation and launch VS Code',
    ]
    for step in vscode_steps:
        doc.add_paragraph(step, style='List Number')
    
    add_screenshot_section(doc, 'VS Code Welcome Screen', 
        'Screenshot showing VS Code after first launch with Welcome tab open')
    
    add_heading_with_style(doc, '3.3 Step 3: Install Essential Extensions', 2)
    doc.add_paragraph('Open Extensions panel (Ctrl+Shift+X) and install:')
    
    extensions = [
        ('Playwright Test for VS Code', 'Run and debug tests directly in VS Code'),
        ('GitHub Copilot', 'AI-powered code suggestions'),
        ('GitHub Copilot Chat', 'Chat interface for AI assistance'),
        ('ESLint', 'Code quality checking'),
        ('Prettier', 'Code formatting'),
    ]
    
    for ext_name, purpose in extensions:
        doc.add_paragraph(f'{ext_name}: {purpose}', style='List Bullet')
    
    add_screenshot_section(doc, 'VS Code Extensions Panel', 
        'Screenshot showing Extensions panel (Ctrl+Shift+X) with Playwright and Copilot extensions highlighted')
    
    add_heading_with_style(doc, '3.4 Step 4: Create Project Directory', 2)
    add_screenshot_section(doc, 'New Folder Creation', 
        'Screenshot of File Explorer creating new folder named "Playwright_Framework"')
    
    add_colored_paragraph(doc, 'Steps:', bold=True, color_rgb=(0, 102, 204))
    steps_4 = [
        'Create folder: C:\\Users\\YourName\\Playwright_Framework',
        'Open VS Code',
        'File → Open Folder → Select the Playwright_Framework folder',
        'VS Code will open with the empty project',
    ]
    for step in steps_4:
        doc.add_paragraph(step, style='List Number')
    
    add_screenshot_section(doc, 'Project Opened in VS Code', 
        'Screenshot showing VS Code with empty Playwright_Framework folder in Explorer panel')
    
    add_heading_with_style(doc, '3.5 Step 5: Open Terminal and Initialize Node Project', 2)
    add_screenshot_section(doc, 'Terminal in VS Code', 
        'Screenshot showing Terminal tab opened at bottom of VS Code (Press Ctrl + backtick)')
    
    add_colored_paragraph(doc, 'Initialize Node project:', bold=True, color_rgb=(0, 102, 204))
    cmd = 'npm init -y'
    doc.add_paragraph(cmd)
    for run in doc.paragraphs[-1].runs:
        run.font.name = 'Courier New'
        run.font.bold = True
    
    doc.add_paragraph('This creates package.json file')
    add_screenshot_section(doc, 'Package.json Created', 
        'Screenshot showing package.json file in project root folder')
    
    doc.add_page_break()
    
    # ===== 4. PHASE 2: MANUAL TEST CASES =====
    add_heading_with_style(doc, '4. Phase 2: Manual Test Case Creation', 1)
    
    doc.add_paragraph(
        'Manual test cases are the foundation of automation. They document the expected behavior of the application '
        'and serve as blueprints for automation scripts.'
    )
    
    add_heading_with_style(doc, '4.1 Test Case Template & Format', 2)
    add_heading_with_style(doc, 'Excel Columns Required:', 3)
    
    tc_template = [
        'TestCaseID - Unique identifier (TC-001)',
        'TestCaseName - Descriptive name of the test',
        'Module - Feature/Module being tested',
        'Description - Detailed description',
        'Priority - High/Medium/Low',
        'Severity - Critical/Major/Minor',
        'Preconditions - Setup requirements',
        'TestData - Input values and test data',
        'Steps - Numbered execution steps',
        'ExpectedResult - Expected outcome',
        'PostConditions - Cleanup steps',
        'Status - Pass/Fail/Not Executed',
    ]
    
    for item in tc_template:
        doc.add_paragraph(item, style='List Bullet')
    
    add_screenshot_section(doc, 'Excel Test Case Sheet', 
        'Screenshot of Excel file with test case columns (TestCaseID, Name, Module, etc.) and sample data rows')
    
    add_heading_with_style(doc, '4.2 Example Manual Test Cases', 2)
    add_colored_paragraph(doc, 'Example 1: User Login Test', bold=True, color_rgb=(0, 102, 204))
    
    tc1_table = doc.add_table(rows=10, cols=2)
    tc1_table.style = 'Light Grid'
    
    tc1_data = [
        ('TestCaseID', 'TC-001'),
        ('TestCaseName', 'Verify valid user login'),
        ('Module', 'Authentication'),
        ('Priority', 'High'),
        ('Preconditions', 'User account exists; User not logged in'),
        ('TestData', 'Username: testuser@example.com; Password: Test@123'),
        ('Steps', '1. Navigate to login page\n2. Enter email\n3. Enter password\n4. Click Login'),
        ('ExpectedResult', 'User logs in successfully; Dashboard displays'),
        ('Status', 'Not Executed'),
    ]
    
    for i, (key, value) in enumerate(tc1_data):
        row = tc1_table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
        shade_cell(row.cells[0], "D9E1F2")
    
    add_screenshot_section(doc, 'Excel Test Case Entry', 
        'Screenshot showing TC-001 filled entry in Excel spreadsheet with all columns')
    
    doc.add_page_break()
    
    # ===== 5. PHASE 3: INSTRUCTIONS.MD =====
    add_heading_with_style(doc, '5. Phase 3: Creating Instructions.md for Copilot', 1)
    
    doc.add_paragraph(
        'The Instructions.md file is crucial - it tells GitHub Copilot Chat how to generate automation scripts. '
        'It contains rules, patterns, best practices, and DemoBlaze application details.'
    )
    
    add_heading_with_style(doc, '5.1 Why Instructions.md Matters', 2)
    
    doc.add_paragraph(
        'GitHub Copilot is an AI model that learns from patterns. By providing clear instructions, rules, and examples, '
        'Copilot generates consistent, high-quality code that follows your framework\'s conventions.'
    )
    
    benefits = [
        'Ensures consistent code quality and style',
        'Guides Copilot to use Page Object Model correctly',
        'Prevents common mistakes in test automation',
        'Reduces manual code review and fixes',
        'Speeds up automation script generation',
        'Maintains framework standards across all tests',
    ]
    for benefit in benefits:
        doc.add_paragraph(benefit, style='List Bullet')
    
    add_heading_with_style(doc, '5.2 Instructions.md File Structure', 2)
    add_colored_paragraph(doc, 'Create file: Instructions_DemoBlaze.md at project root', bold=True, color_rgb=(0, 102, 204))
    add_screenshot_section(doc, 'Instructions File Created', 
        'Screenshot showing Instructions_DemoBlaze.md file in VS Code project root in Explorer')
    
    add_heading_with_style(doc, '5.3 Key Sections in Instructions.md', 2)
    
    sections = {
        'Quick Start': 'Explains how user provides test case info',
        'Target Application Details': 'URL, features, and capabilities of the app being tested',
        'Test Case Source Format': 'Expected Excel format and columns',
        'Critical Rules': 'Mandatory rules for Copilot to follow',
        'Rule Examples': 'Before/After code examples showing correct patterns',
        'Application-Specific Selectors': 'DemoBlaze page locators and element IDs',
        'Test Naming Conventions': 'How to name test methods and classes',
    }
    
    for section, purpose in sections.items():
        doc.add_paragraph(f'{section}: {purpose}', style='List Bullet')
    
    add_heading_with_style(doc, '5.4 Critical Rules to Include', 2)
    
    critical_rules = [
        'Rule 1: Always Display Test Case First AND Wait for Confirmation',
        'Rule 2: Capture BOTH Action AND Validation Locators',
        'Rule 3: Never Guess Locators — Verify First',
        'Rule 4: Use Locator Priority (data-testid > id > class > XPath)',
        'Rule 5: Use Page Object Model with Proper Structure',
        'Rule 6: Element Names Must Use "Locator" Suffix',
        'Rule 7: Avoid Complex Assertions in Test Methods',
        'Rule 8: No Hard Waits — Use Playwright Built-in Waits',
        'Rule 9: Implement ALL Test Steps — No Partial Implementation',
        'Rule 10: Follow Test Method Naming Convention',
    ]
    
    for rule in critical_rules:
        doc.add_paragraph(rule, style='List Bullet')
    
    add_screenshot_section(doc, 'Instructions.md Content', 
        'Screenshot showing Instructions_DemoBlaze.md file open in VS Code with all 10 critical rules visible')
    
    doc.add_page_break()
    
    # ===== 6. PHASE 4: COPILOT CHAT =====
    add_heading_with_style(doc, '6. Phase 4: Using Copilot Chat to Automate Tests', 1)
    
    doc.add_paragraph(
        'This is where the magic happens. GitHub Copilot Chat uses the Instructions.md file to generate '
        'high-quality, framework-compliant test automation code.'
    )
    
    add_heading_with_style(doc, '6.1 Step 1: Open Copilot Chat', 2)
    
    steps_chat = [
        'Open Copilot Chat in VS Code (Ctrl+Shift+I or View → Copilot Chat)',
        'Click the paperclip icon to attach Instructions_DemoBlaze.md',
        'Or use @ symbol: type @Instructions_DemoBlaze.md',
        'Copilot will read and understand the instructions',
    ]
    
    for step in steps_chat:
        doc.add_paragraph(step, style='List Number')
    
    add_screenshot_section(doc, 'Copilot Chat Interface', 
        'Screenshot of Copilot Chat panel opened with @ mention for Instructions.md file')
    
    add_heading_with_style(doc, '6.2 Step 2: Provide Test Case Details', 2)
    
    doc.add_paragraph('Format your request to Copilot as:')
    
    sample_prompt = '''Using Instructions_DemoBlaze.md, automate TC-001:
- TestCaseID: TC-001
- TestCaseName: Verify valid user login
- TestData: Username: testuser@example.com, Password: Test@123
- Steps: Navigate → Enter email → Enter password → Click Login
- ExpectedResult: User logs in successfully, Dashboard displays

Provide loginPage.ts and auth.spec.ts files.'''
    
    doc.add_paragraph(sample_prompt, style='No Spacing')
    for run in doc.paragraphs[-1].runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    add_screenshot_section(doc, 'Copilot Chat Prompt', 
        'Screenshot showing test case details entered in Copilot Chat message box')
    
    add_heading_with_style(doc, '6.3 Step 3: Copilot Generates Code', 2)
    
    doc.add_paragraph('Copilot will generate:')
    
    generated_items = [
        'Login Page Object (loginPage.ts) with locators and methods',
        'Test Specification (auth.spec.ts) with test cases',
        'Complete TypeScript code following all Instructions.md rules',
    ]
    
    for item in generated_items:
        doc.add_paragraph(item, style='List Bullet')
    
    add_screenshot_section(doc, 'Copilot Generated Code', 
        'Screenshot showing Copilot Chat response with generated Page Object TypeScript code')
    
    add_heading_with_style(doc, '6.4 Step 4: Copy Code to Files', 2)
    
    copy_steps = [
        'Copy the Page Object code from chat',
        'Create Pages/loginPage.ts and paste code',
        'Copy the Test Specification code',
        'Create tests/auth.spec.ts and paste code',
    ]
    
    for step in copy_steps:
        doc.add_paragraph(step, style='List Number')
    
    add_screenshot_section(doc, 'Files Created in VS Code', 
        'Screenshot showing Pages/loginPage.ts and tests/auth.spec.ts files created with code pasted')
    
    doc.add_page_break()
    
    # ===== 7. PROJECT STRUCTURE =====
    add_heading_with_style(doc, '7. Phase 5: Project Structure & Organization', 1)
    
    add_heading_with_style(doc, '7.1 Complete Folder Structure', 2)
    
    structure_text = '''Playwright_Framework/
├── Pages/
│   ├── loginPage.ts
│   ├── homePage.ts
│   └── cartPage.ts
├── tests/
│   ├── auth.spec.ts
│   ├── shopping.spec.ts
│   └── checkout.spec.ts
├── utilities/
│   ├── testData.ts
│   └── helpers.ts
├── test-cases_manual/
│   └── TestCases.xlsx
├── Instructions_DemoBlaze.md
├── playwright.config.ts
├── package.json
└── README.md'''
    
    doc.add_paragraph(structure_text, style='No Spacing')
    for run in doc.paragraphs[-1].runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    add_screenshot_section(doc, 'Project Explorer Structure', 
        'Screenshot of VS Code Explorer showing complete folder structure with all folders and files')
    
    doc.add_page_break()
    
    # ===== 8. WORKFLOW =====
    add_heading_with_style(doc, '8. Phase 6: Complete Automation Workflow', 1)
    
    workflow_steps = [
        ('Manual Test Case', 'Document test in Excel with all details'),
        ('Review & Approve', 'Stakeholders review the manual test case'),
        ('Copilot Request', 'Use Copilot Chat with Instructions.md'),
        ('Code Generated', 'Copilot generates Page Object and Test files'),
        ('Code Review', 'Review generated code, verify locators'),
        ('Execute Test', 'Run test and verify it works'),
        ('Update Status', 'Mark test as "Automated" in Excel'),
    ]
    
    for i, (step_name, step_desc) in enumerate(workflow_steps):
        doc.add_paragraph(f'{i+1}. {step_name}: {step_desc}', style='List Bullet')
    
    add_screenshot_section(doc, 'Complete Automation Workflow', 
        'Screenshot or flowchart showing: Excel TC → Copilot Chat → Generated Code → Test Execution → Report')
    
    doc.add_page_break()
    
    # ===== 9. COPILOT RULES =====
    add_heading_with_style(doc, '9. Copilot Rules & Best Practices', 1)
    
    add_heading_with_style(doc, '9.1 The 10 Critical Rules Summary', 2)
    
    rules = {
        'Rule 1: Always Display Test Case First': 'Display complete test case, WAIT for "Yes, proceed" before coding',
        'Rule 2: Capture Both Locators': 'Capture action locators (buttons) AND validation locators (messages)',
        'Rule 3: Never Guess Locators': 'Inspect actual DOM, verify selector works before using',
        'Rule 4: Use Locator Priority': 'Prefer data-testid > id > class > CSS > XPath',
        'Rule 5: Use Page Object Model': 'Declare elements as properties, use methods for interactions',
        'Rule 6: Use Locator Suffix': 'Element names: loginButtonLocator (not loginButton)',
        'Rule 7: Avoid Complex Assertions': 'Create assertion methods in page objects, keep tests clean',
        'Rule 8: No Hard Waits': 'Use waitFor() with state conditions instead of setTimeout()',
        'Rule 9: Implement ALL Steps': 'Complete all test steps and validations, no partial implementations',
        'Rule 10: Follow Naming': 'Pattern: demoBlaze_[page]_[description]() like demoBlaze_login_verifyValidUserLogin()',
    }
    
    for rule_title, rule_desc in rules.items():
        doc.add_paragraph(f'{rule_title}:', style='List Bullet')
        doc.add_paragraph(rule_desc, style='List Bullet 2')
    
    doc.add_page_break()
    
    # ===== 10. EXAMPLE =====
    add_heading_with_style(doc, '10. Step-by-Step Example: Add to Cart', 1)
    
    add_colored_paragraph(doc, 'Step 1: Create Manual Test Case in Excel', bold=True, color_rgb=(0, 102, 204))
    
    manual_tc = '''TestCaseID: TC-002
TestCaseName: Verify user can add single product to cart
Module: Shopping Cart
Preconditions: User is on DemoBlaze home page
TestData: Product: Samsung Galaxy S10
Steps:
  1. Navigate to https://demoblaze.com/index.html
  2. Look for Samsung Galaxy S10 product
  3. Click "Add to cart" button
  4. Verify success message appears
ExpectedResult: Product added to cart, Cart count increased'''
    
    doc.add_paragraph(manual_tc, style='No Spacing')
    for run in doc.paragraphs[-1].runs:
        run.font.size = Pt(9)
    
    add_screenshot_section(doc, 'Manual TC in Excel', 
        'Screenshot of TC-002 entry in Excel with all columns: TestCaseID, Name, Module, Steps, ExpectedResult')
    
    add_colored_paragraph(doc, 'Step 2: Request Copilot to Automate', bold=True, color_rgb=(0, 102, 204))
    
    copilot_msg = '''@Instructions_DemoBlaze.md
Please automate TC-002: Add to Cart
Generate homePage.ts and shopping.spec.ts files.'''
    
    doc.add_paragraph(copilot_msg, style='No Spacing')
    for run in doc.paragraphs[-1].runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    add_screenshot_section(doc, 'Request Sent to Copilot', 
        'Screenshot showing TC-002 automation request sent in Copilot Chat')
    
    add_colored_paragraph(doc, 'Step 3: Execute Generated Test', bold=True, color_rgb=(0, 102, 204))
    
    exec_cmd = 'npx playwright test tests/shopping.spec.ts --headed'
    cmd_para = doc.add_paragraph(exec_cmd)
    for run in cmd_para.runs:
        run.font.name = 'Courier New'
        run.font.bold = True
    
    add_screenshot_section(doc, 'Test Execution in Browser', 
        'Screenshot showing browser window with DemoBlaze application running the automated test')
    
    add_colored_paragraph(doc, 'Step 4: View Test Report', bold=True, color_rgb=(0, 102, 204))
    
    report_cmd = 'npx playwright show-report'
    report_para = doc.add_paragraph(report_cmd)
    for run in report_para.runs:
        run.font.name = 'Courier New'
        run.font.bold = True
    
    add_screenshot_section(doc, 'Test Report', 
        'Screenshot of HTML report showing TC-002 passed with execution time, screenshots, and video')
    
    doc.add_page_break()
    
    # ===== 11. TESTING & EXECUTION =====
    add_heading_with_style(doc, '11. Testing & Execution', 1)
    
    add_heading_with_style(doc, '11.1 Essential Test Commands', 2)
    
    test_commands = [
        ('All Tests', 'npx playwright test'),
        ('Specific File', 'npx playwright test tests/auth.spec.ts'),
        ('Specific Test', 'npx playwright test -g "verify valid user"'),
        ('Headed Mode', 'npx playwright test --headed'),
        ('Debug Mode', 'npx playwright test --debug'),
        ('Chromium Only', 'npx playwright test --project=chromium'),
        ('Verbose', 'npx playwright test --verbose'),
    ]
    
    for cmd_name, cmd in test_commands:
        doc.add_paragraph(f'{cmd_name}:')
        cmd_para = doc.add_paragraph(cmd)
        for run in cmd_para.runs:
            run.font.name = 'Courier New'
            run.font.bold = True
    
    add_screenshot_section(doc, 'Terminal Test Execution', 
        'Screenshot showing test execution in terminal with progress, passing tests, and duration')
    
    add_heading_with_style(doc, '11.2 Viewing Reports', 2)
    
    doc.add_paragraph('View the generated HTML report:')
    report_para = doc.add_paragraph('npx playwright show-report')
    for run in report_para.runs:
        run.font.name = 'Courier New'
        run.font.bold = True
    
    add_screenshot_section(doc, 'HTML Test Report Interface', 
        'Screenshot of HTML report showing test results summary, passed count, failed count, and detailed test logs')
    
    doc.add_page_break()
    
    # ===== 12. TROUBLESHOOTING =====
    add_heading_with_style(doc, '12. Troubleshooting & Tips', 1)
    
    add_heading_with_style(doc, '12.1 Common Issues and Solutions', 2)
    
    issues_table = doc.add_table(rows=8, cols=2)
    issues_table.style = 'Light Grid Accent 1'
    
    header_cells = issues_table.rows[0].cells
    header_cells[0].text = 'Issue'
    header_cells[1].text = 'Solution'
    shade_cell(header_cells[0], "4472C4")
    shade_cell(header_cells[1], "4472C4")
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    issues = [
        ('Element not found', 'Check locator is correct using page.screenshot(), check DOM'),
        ('Test timeout', 'Increase timeout in playwright.config.ts, add waitFor()'),
        ('Flaky tests', 'Use waitFor() with state conditions, avoid setTimeout()'),
        ('Copilot code incorrect', 'Provide detailed Instructions.md, ask Copilot to refactor'),
        ('Module not found', 'Run npm install, check import paths in files'),
        ('Tests slow', 'Enable parallelization, optimize waits, reduce delays'),
        ('Report not generating', 'Verify HTML reporter in config, check file permissions'),
    ]
    
    for i, (issue, solution) in enumerate(issues):
        row = issues_table.rows[i + 1]
        row.cells[0].text = issue
        row.cells[1].text = solution
    
    add_heading_with_style(doc, '12.2 Tips for Success', 2)
    
    tips = [
        'Keep Instructions.md updated with new patterns and application selectors',
        'Review Copilot-generated code before using it in production',
        'Use meaningful variable names ending with "Locator"',
        'Test generated code in headed mode first (--headed flag)',
        'Maintain test data separately in utilities/testData.ts',
        'Use Page Object Model consistently across all tests',
        'Add comments to complex test logic',
        'Run tests regularly to catch failures early',
        'Update locators when application UI changes',
        'Use Git version control for all test code',
    ]
    
    for tip in tips:
        doc.add_paragraph(tip, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== CONCLUSION =====
    add_heading_with_style(doc, 'Conclusion', 1)
    
    doc.add_paragraph(
        'You now have a complete guide for building and maintaining a Playwright test automation framework '
        'with GitHub Copilot Chat assistance. This methodology enables you to:'
    )
    
    conclusion_points = [
        'Create well-structured manual test cases in Excel',
        'Use Instructions.md to guide Copilot in code generation',
        'Generate consistent, maintainable automation scripts',
        'Execute tests and view comprehensive HTML reports',
        'Scale your test suite efficiently',
        'Maintain code quality and framework standards',
    ]
    
    for point in conclusion_points:
        doc.add_paragraph(point, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('Quick Start Checklist:')
    checklist = [
        '☐ Install Node.js and VS Code',
        '☐ Install Playwright and dependencies',
        '☐ Create project structure (Pages, tests, utilities)',
        '☐ Create Instructions_DemoBlaze.md with 10 critical rules',
        '☐ Create first manual test case in Excel',
        '☐ Request Copilot Chat to generate automation code',
        '☐ Copy generated code to Page Object and Test files',
        '☐ Run: npx playwright test',
        '☐ View report: npx playwright show-report',
        '☐ Repeat for more test cases',
    ]
    
    for item in checklist:
        doc.add_paragraph(item, style='List Bullet')
    
    # Save document
    output_path = r'd:\My Data\Automation-playwright\Playwright_Framework\Playwright_Framework_FINAL_WITH_IMAGES.docx'
    doc.save(output_path)
    print(f"Final document with placeholder images generated: {output_path}")
    return output_path

if __name__ == "__main__":
    create_final_document_with_images()
