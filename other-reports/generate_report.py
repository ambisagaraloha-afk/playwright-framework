from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

def add_heading_with_style(doc, text, level, color_rgb=(0, 51, 102)):
    """Add heading with custom styling"""
    heading = doc.add_heading(text, level=level)
    heading.runs[0].font.color.rgb = RGBColor(*color_rgb)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_colored_paragraph(doc, text, color_rgb=(0, 0, 0), bold=False, size=11):
    """Add paragraph with custom color"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(*color_rgb)
    run.font.bold = bold
    run.font.size = Pt(size)
    return p

def shade_cell(cell, fill_color):
    """Shade a table cell with color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_comprehensive_report():
    """Create a comprehensive Playwright framework report"""
    
    doc = Document()
    
    # Set up document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title Page
    title = doc.add_heading('Playwright E2E Automation Framework', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph('DemoBlaze Application Testing')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_paragraph()  # Spacing
    
    # Project Details
    details_table = doc.add_table(rows=6, cols=2)
    details_table.style = 'Light Grid Accent 1'
    
    details = [
        ('Project Name', 'Playwright Framework - DemoBlaze'),
        ('Framework Type', 'Page Object Model (POM)'),
        ('Technology Stack', 'Playwright + TypeScript'),
        ('Test Application', 'https://demoblaze.com/index.html'),
        ('Report Generated', datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ('Status', 'Active Development')
    ]
    
    for i, (key, value) in enumerate(details):
        key_cell = details_table.rows[i].cells[0]
        val_cell = details_table.rows[i].cells[1]
        key_cell.text = key
        val_cell.text = value
        shade_cell(key_cell, "D3D3D3")
    
    doc.add_page_break()
    
    # Table of Contents
    add_heading_with_style(doc, 'Table of Contents', 1)
    toc_items = [
        '1. Project Overview',
        '2. Setup & Installation',
        '   2.1 Prerequisites',
        '   2.2 Step-by-Step Installation',
        '3. Framework Architecture',
        '4. Folder Structure',
        '5. Technology Stack',
        '6. Manual Test Cases',
        '7. Automation Scripts',
        '8. Running Tests',
        '9. Test Reports',
        '10. Best Practices',
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 1. Project Overview
    add_heading_with_style(doc, '1. Project Overview', 1)
    doc.add_paragraph(
        'This Playwright framework is designed for end-to-end testing of the DemoBlaze e-commerce application. '
        'It follows the Page Object Model (POM) architecture, ensuring maintainability, scalability, and reusability of test code.'
    )
    
    add_heading_with_style(doc, '1.1 Framework Objectives', 2)
    objectives = [
        'Automate critical user journeys on DemoBlaze application',
        'Ensure consistent test execution across multiple browsers',
        'Maintain clean, maintainable, and scalable test code',
        'Generate detailed HTML test reports',
        'Implement robust error handling and logging',
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
    
    add_heading_with_style(doc, '1.2 Test Application Overview', 2)
    app_features = [
        'Product Categories: Phones, Laptops, Monitors',
        'User Authentication: Sign up, Login, Logout',
        'Shopping Cart: Add/Remove products, View cart',
        'Checkout: Order placement with payment details',
        'UI Features: Product carousel, search, filtering, responsive design',
    ]
    for feature in app_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # 2. Setup & Installation
    add_heading_with_style(doc, '2. Setup & Installation', 1)
    
    add_heading_with_style(doc, '2.1 Prerequisites', 2)
    
    prerequisites_table = doc.add_table(rows=5, cols=3)
    prerequisites_table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = prerequisites_table.rows[0].cells
    headers = ['Software', 'Version', 'Download Link']
    for i, header_text in enumerate(headers):
        header_cells[i].text = header_text
        shade_cell(header_cells[i], "4472C4")
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    prereqs = [
        ('Node.js', '18.x or higher', 'https://nodejs.org/'),
        ('Visual Studio Code', 'Latest version', 'https://code.visualstudio.com/'),
        ('Git', 'Latest version', 'https://git-scm.com/'),
        ('TypeScript', '5.x or higher', 'npm install -g typescript'),
    ]
    
    for i, (software, version, link) in enumerate(prereqs):
        row = prerequisites_table.rows[i + 1]
        row.cells[0].text = software
        row.cells[1].text = version
        row.cells[2].text = link
    
    add_heading_with_style(doc, '2.2 Step-by-Step Installation', 2)
    
    steps = [
        ('Step 1: Install Node.js', [
            'Download Node.js from https://nodejs.org/',
            'Run the installer and follow the setup wizard',
            'Verify installation: Open Command Prompt and type: node --version',
            'Check npm: npm --version',
        ]),
        ('Step 2: Install Visual Studio Code', [
            'Download from https://code.visualstudio.com/',
            'Run installer and complete the installation',
            'Launch VS Code',
        ]),
        ('Step 3: Install Essential VS Code Extensions', [
            'Playwright Test for VS Code',
            'GitHub Copilot (for AI assistance)',
            'TypeScript Vue Plugin',
            'Prettier - Code formatter',
            'REST Client',
        ]),
        ('Step 4: Create Project Directory', [
            'Create a new folder: Playwright_Framework',
            'Open the folder in VS Code',
            'Open terminal in VS Code (Ctrl + `)',
        ]),
        ('Step 5: Initialize Node Project', [
            'Run command: npm init -y',
            'This creates package.json file',
        ]),
        ('Step 6: Install Playwright', [
            'Run command: npm install --save-dev @playwright/test',
            'Run command: npm install --save-dev @types/node',
            'Install xlsx for Excel reading: npm install xlsx',
        ]),
        ('Step 7: Create Project Structure', [
            'Create folders: Pages, tests, utilities, test-cases_manual',
            'Create playwright.config.ts file',
            'Configure browser settings and reporters',
        ]),
    ]
    
    for step_title, step_items in steps:
        add_heading_with_style(doc, step_title, 3, (0, 102, 204))
        for item in step_items:
            doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 3. Framework Architecture
    add_heading_with_style(doc, '3. Framework Architecture', 1)
    
    doc.add_paragraph(
        'The framework follows the Page Object Model (POM) design pattern, which separates test logic from locator information. '
        'This approach improves maintainability and reduces code duplication.'
    )
    
    add_heading_with_style(doc, '3.1 Design Pattern: Page Object Model (POM)', 2)
    doc.add_paragraph(
        'POM is a design pattern that allows you to create an abstraction of the tested web application pages. '
        'Each page is represented as a class, and the UI elements of that page are defined as variables of the class. '
        'The methods of the class represent the possible actions that can be performed on a page.'
    )
    
    benefits = [
        'Improved code maintainability',
        'Reduced code duplication',
        'Easy to update locators in one place',
        'Better test readability',
        'Scalable test framework',
    ]
    
    for benefit in benefits:
        doc.add_paragraph(benefit, style='List Bullet')
    
    doc.add_page_break()
    
    # 4. Folder Structure
    add_heading_with_style(doc, '4. Folder Structure', 1)
    
    structure_text = '''
Playwright_Framework/
├── Pages/                          # Page Object Model classes
│   ├── homePage.ts                # Home page locators and methods
│   ├── loginPage.ts               # Login page locators and methods
│   └── catPage.ts                 # Category page locators and methods
│
├── tests/                          # Test specifications
│   ├── loginTest.spec.ts          # Login test cases
│   ├── addToCart.spec.ts          # Add to cart test cases
│   ├── filterByCategory.spec.ts   # Category filtering test cases
│   └── example.spec.ts            # Example test template
│
├── utilities/                      # Utility files
│   ├── testData.ts                # Test data and credentials
│   └── helpers.ts                 # Helper functions (if added)
│
├── test-cases_manual/             # Manual test cases
│   ├── DemoBlaze_TestCases.xlsx   # Test cases in Excel format
│   └── generate_test_cases_excel.js  # Excel generation utility
│
├── playwright-report/             # Test execution reports
│   └── index.html                 # HTML report
│
├── playwright.config.ts           # Playwright configuration
├── package.json                   # Project dependencies
├── Instructions_DemoBlaze.md      # Copilot instructions
└── README.md                      # Project documentation
    '''
    
    doc.add_paragraph(structure_text, style='No Spacing')
    
    add_heading_with_style(doc, '4.1 Folder Descriptions', 2)
    
    folders = {
        'Pages/': 'Contains Page Object Model classes. Each class represents a page in the application and holds locators and page interaction methods.',
        'tests/': 'Contains Playwright test specification files. Each file contains test cases for a specific feature or module.',
        'utilities/': 'Contains utility files like test data, helper functions, and configuration constants.',
        'test-cases_manual/': 'Contains manual test cases in Excel format before automation.',
        'playwright-report/': 'Auto-generated folder containing HTML test reports.',
    }
    
    for folder, description in folders.items():
        add_colored_paragraph(doc, f'{folder}', bold=True, color_rgb=(0, 102, 204))
        doc.add_paragraph(description)
    
    doc.add_page_break()
    
    # 5. Technology Stack
    add_heading_with_style(doc, '5. Technology Stack', 1)
    
    tech_table = doc.add_table(rows=8, cols=3)
    tech_table.style = 'Light Grid Accent 1'
    
    # Header
    header_cells = tech_table.rows[0].cells
    tech_headers = ['Technology', 'Version', 'Purpose']
    for i, header_text in enumerate(tech_headers):
        header_cells[i].text = header_text
        shade_cell(header_cells[i], "4472C4")
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    techs = [
        ('Playwright', '1.56.1', 'End-to-End testing framework'),
        ('TypeScript', '5.x', 'Type-safe JavaScript'),
        ('Node.js', '18.x+', 'JavaScript runtime'),
        ('xlsx', '0.18.5', 'Excel file parsing'),
        ('VS Code', 'Latest', 'Code editor'),
        ('@types/node', 'Latest', 'TypeScript definitions for Node.js'),
        ('GitHub Copilot', 'Latest', 'AI-powered code assistance'),
    ]
    
    for i, (tech, version, purpose) in enumerate(techs):
        row = tech_table.rows[i + 1]
        row.cells[0].text = tech
        row.cells[1].text = version
        row.cells[2].text = purpose
    
    doc.add_page_break()
    
    # 6. Manual Test Cases
    add_heading_with_style(doc, '6. Manual Test Cases', 1)
    
    doc.add_paragraph(
        'Before automating, manual test cases are documented in an Excel format. '
        'These serve as the blueprint for creating automation scripts.'
    )
    
    add_heading_with_style(doc, '6.1 Excel Test Case Format', 2)
    
    excel_cols = {
        'TestCaseId': 'Unique identifier (e.g., TC-001)',
        'TestCaseName': 'Descriptive name of the test',
        'Description': 'Detailed description of what is being tested',
        'Category': 'Test category (UI/Navigation, Authentication, Shopping Cart, etc.)',
        'Priority': 'High / Medium / Low',
        'Preconditions': 'Setup requirements before test execution',
        'TestData': 'Input values needed for the test',
        'Steps': 'Numbered steps to perform the test',
        'ExpectedResult': 'Expected outcome of the test',
        'Status': 'Test execution status',
    }
    
    for col, description in excel_cols.items():
        add_colored_paragraph(doc, f'{col}:', bold=True, color_rgb=(0, 102, 204))
        doc.add_paragraph(description)
    
    add_heading_with_style(doc, '6.2 Sample Test Case', 2)
    
    sample_table = doc.add_table(rows=11, cols=2)
    sample_table.style = 'Light Grid Accent 1'
    
    sample_data = [
        ('Test Case ID', 'TC-001'),
        ('Test Case Name', 'Verify Home Page Loads Successfully'),
        ('Category', 'UI/Navigation'),
        ('Priority', 'High'),
        ('Preconditions', 'User has internet connectivity, Access to demoblaze.com'),
        ('Test Data', 'URL: https://demoblaze.com/index.html'),
        ('Steps', '1. Navigate to URL\\n2. Wait for page to load\\n3. Verify all elements visible'),
        ('Expected Result', 'Home page loads with all categories and products visible'),
        ('Status', 'Not Executed'),
    ]
    
    for i, (key, value) in enumerate(sample_data):
        row = sample_table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
        shade_cell(row.cells[0], "E7E6E6")
    
    doc.add_page_break()
    
    # 7. Automation Scripts
    add_heading_with_style(doc, '7. Automation Scripts', 1)
    
    add_heading_with_style(doc, '7.1 Page Object Model Example - Home Page', 2)
    
    doc.add_paragraph('File: Pages/homePage.ts')
    
    home_page_code = '''import { Locator, Page } from "@playwright/test";

export class Homepage {
    readonly page: Page;
    readonly phonesCategoryLocator: Locator;
    readonly laptopsCategoryLocator: Locator;
    readonly monitorsCategoryLocator: Locator;
    readonly productItemsLocator: Locator;

    constructor(page: Page) {
        this.page = page;
        this.phonesCategoryLocator = page.locator('//a[contains(., "Phones")]');
        this.laptopsCategoryLocator = page.locator('//a[contains(., "Laptops")]');
        this.monitorsCategoryLocator = page.locator('//a[contains(., "Monitors")]');
        this.productItemsLocator = page.locator('//div[@class="col-lg-4 col-md-6 mb-4"]');
    }

    async filterByPhonesCategory() {
        await this.phonesCategoryLocator.waitFor({ state: 'visible', timeout: 5000 });
        await this.phonesCategoryLocator.click();
        await this.page.waitForTimeout(1500);
    }

    async filterByLaptopsCategory() {
        await this.laptopsCategoryLocator.waitFor({ state: 'visible', timeout: 5000 });
        await this.laptopsCategoryLocator.click();
        await this.page.waitForTimeout(1500);
    }

    async filterByMonitorsCategory() {
        await this.monitorsCategoryLocator.waitFor({ state: 'visible', timeout: 5000 });
        await this.monitorsCategoryLocator.click();
        await this.page.waitForTimeout(1500);
    }
}'''
    
    code_para = doc.add_paragraph(home_page_code, style='No Spacing')
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    add_heading_with_style(doc, '7.2 Test Specification Example - Login Test', 2)
    
    doc.add_paragraph('File: tests/loginTest.spec.ts')
    
    login_test_code = '''import { test } from '@playwright/test';
import { LoginPage } from '../Pages/loginPage';
import { credentials } from '../utilities/testData';

test.beforeEach(async ({ page }) => {
    await page.goto("https://demoblaze.com/index.html");
});

test("login page test case", async ({ page }) => {
    const loginPageObj = new LoginPage(page);
    await loginPageObj.login(credentials.userName, credentials.password);
});'''
    
    code_para = doc.add_paragraph(login_test_code, style='No Spacing')
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # 8. Running Tests
    add_heading_with_style(doc, '8. Running Tests', 1)
    
    add_heading_with_style(doc, '8.1 Run All Tests', 2)
    run_all_cmd = 'npm run test  OR  npx playwright test'
    doc.add_paragraph(run_all_cmd, style='No Spacing')
    
    add_heading_with_style(doc, '8.2 Run Specific Test File', 2)
    run_specific_cmd = 'npx playwright test tests/loginTest.spec.ts'
    doc.add_paragraph(run_specific_cmd, style='No Spacing')
    
    add_heading_with_style(doc, '8.3 Run Tests in Specific Browser', 2)
    run_browser_cmds = [
        'npx playwright test --project=chromium',
        'npx playwright test --project=firefox',
        'npx playwright test --project=webkit',
    ]
    for cmd in run_browser_cmds:
        doc.add_paragraph(cmd, style='List Bullet')
    
    add_heading_with_style(doc, '8.4 Run Tests in Headed Mode', 2)
    run_headed_cmd = 'npx playwright test --headed'
    doc.add_paragraph(run_headed_cmd, style='No Spacing')
    doc.add_paragraph('This shows the browser window during test execution.')
    
    add_heading_with_style(doc, '8.5 Run Tests with Debug Mode', 2)
    run_debug_cmd = 'npx playwright test --debug'
    doc.add_paragraph(run_debug_cmd, style='No Spacing')
    doc.add_paragraph('This opens Playwright Inspector for debugging.')
    
    doc.add_page_break()
    
    # 9. Test Reports
    add_heading_with_style(doc, '9. Test Reports', 1)
    
    add_heading_with_style(doc, '9.1 Generate HTML Report', 2)
    doc.add_paragraph(
        'Playwright automatically generates an HTML report after test execution. '
        'The report is stored in the playwright-report/ folder.'
    )
    
    add_heading_with_style(doc, '9.2 View HTML Report', 2)
    view_report_cmd = 'npx playwright show-report'
    doc.add_paragraph(view_report_cmd, style='No Spacing')
    doc.add_paragraph('This opens the HTML report in your default browser.')
    
    add_heading_with_style(doc, '9.3 Report Features', 2)
    report_features = [
        'Test execution summary (passed, failed, skipped)',
        'Detailed test logs and screenshots',
        'Video recordings of failed tests',
        'Trace files for debugging',
        'Timing and performance metrics',
        'Browser compatibility information',
    ]
    for feature in report_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # 10. Best Practices
    add_heading_with_style(doc, '10. Best Practices', 1)
    
    best_practices = {
        'Use Page Object Model': [
            'Separate locators from test logic',
            'Create reusable page methods',
            'Maintain consistency across tests',
        ],
        'Explicit Waits': [
            'Use waitFor() instead of hardcoded delays',
            'Set appropriate timeouts',
            'Wait for specific element conditions',
        ],
        'Error Handling': [
            'Use try-catch blocks for critical operations',
            'Provide meaningful error messages',
            'Log important test steps',
        ],
        'Test Data Management': [
            'Store test data in separate files',
            'Use constants for repeated values',
            'Avoid hardcoding credentials',
        ],
        'Code Quality': [
            'Follow TypeScript best practices',
            'Use meaningful variable and method names',
            'Add comments for complex logic',
            'Keep tests focused and isolated',
        ],
        'CI/CD Integration': [
            'Use headless mode for CI pipelines',
            'Configure parallel test execution',
            'Archive test reports automatically',
        ],
    }
    
    for practice, tips in best_practices.items():
        add_heading_with_style(doc, practice, 3, (0, 102, 204))
        for tip in tips:
            doc.add_paragraph(tip, style='List Bullet')
    
    doc.add_page_break()
    
    # Conclusion
    add_heading_with_style(doc, 'Conclusion', 1)
    doc.add_paragraph(
        'This Playwright framework provides a robust foundation for end-to-end testing of web applications. '
        'By following the Page Object Model pattern and the best practices outlined in this document, '
        'you can create scalable, maintainable, and efficient test automation suites.'
    )
    doc.add_paragraph()
    doc.add_paragraph('Key benefits of this framework:')
    benefits = [
        'Easy to maintain and update locators',
        'Reusable page components across tests',
        'Comprehensive test reporting',
        'Cross-browser testing capabilities',
        'CI/CD integration ready',
        'TypeScript type safety',
    ]
    for benefit in benefits:
        doc.add_paragraph(benefit, style='List Bullet')
    
    # Save document
    output_path = r'd:\My Data\Automation-playwright\Playwright_Framework\Playwright_Framework_Report.docx'
    doc.save(output_path)
    print(f"Report generated successfully: {output_path}")

if __name__ == "__main__":
    create_comprehensive_report()
